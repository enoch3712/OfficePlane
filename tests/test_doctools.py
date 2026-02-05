"""
Tests for the doctools module.

Tests cover:
1. Result types (Ok/Err pattern)
2. DocumentEditor (batch operations, transactions)
3. High-level operations (StructureReader, ContentModifier, TableBuilder)
4. Planning system (DocumentPlan, PlanExecutor)
"""

import tempfile
import shutil
from pathlib import Path

import pytest
from docx import Document

from officeplane.doctools import (
    # Result types
    Result,
    Ok,
    Err,
    DocError,
    ErrorCode,
    # Editor
    DocumentEditor,
    EditSession,
    # Operations
    StructureReader,
    ContentModifier,
    TableBuilder,
    # Planner
    DocumentPlan,
    PlanPhase,
    ActionStep,
    PlanExecutor,
)
from officeplane.doctools.result import collect_results, try_operation
from officeplane.doctools.planner import ActionType, plan_from_dict


# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    path = Path(tempfile.mkdtemp())
    yield path
    shutil.rmtree(path)


@pytest.fixture
def empty_doc(temp_dir):
    """Create an empty document."""
    path = temp_dir / "empty.docx"
    doc = Document()
    doc.save(str(path))
    return path


@pytest.fixture
def sample_doc(temp_dir):
    """Create a sample document with structure."""
    path = temp_dir / "sample.docx"
    doc = Document()

    # Add title
    doc.add_heading("Sample Document", level=0)
    doc.add_paragraph("This is the introduction paragraph.")

    # Chapter 1
    doc.add_heading("Chapter 1: Getting Started", level=1)
    doc.add_paragraph("Welcome to chapter one.")
    doc.add_paragraph("This chapter covers the basics.")

    # Section 1.1
    doc.add_heading("Section 1.1: Prerequisites", level=2)
    doc.add_paragraph("You will need Python 3.9+")

    # Chapter 2
    doc.add_heading("Chapter 2: Advanced Topics", level=1)
    doc.add_paragraph("Now we dive deeper.")

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "100"
    table.cell(2, 0).text = "Beta"
    table.cell(2, 1).text = "200"

    doc.save(str(path))
    return path


# =============================================================================
# Test Result Types
# =============================================================================


class TestResultTypes:
    """Tests for Result, Ok, Err, DocError."""

    def test_ok_is_ok(self):
        result: Result[int] = Ok(42)
        assert result.is_ok()
        assert not result.is_err()

    def test_ok_unwrap(self):
        result: Result[str] = Ok("hello")
        assert result.unwrap() == "hello"

    def test_ok_unwrap_or(self):
        result: Result[int] = Ok(42)
        assert result.unwrap_or(0) == 42

    def test_err_is_err(self):
        error = DocError(code=ErrorCode.FILE_NOT_FOUND, message="Not found")
        result: Result[int] = Err(error)
        assert result.is_err()
        assert not result.is_ok()

    def test_err_unwrap_raises(self):
        error = DocError(code=ErrorCode.FILE_NOT_FOUND, message="Not found")
        result: Result[int] = Err(error)
        with pytest.raises(ValueError):
            result.unwrap()

    def test_err_unwrap_or(self):
        error = DocError(code=ErrorCode.FILE_NOT_FOUND, message="Not found")
        result: Result[int] = Err(error)
        assert result.unwrap_or(99) == 99

    def test_ok_map(self):
        result: Result[int] = Ok(10)
        mapped = result.map(lambda x: x * 2)
        assert mapped.is_ok()
        assert mapped.unwrap() == 20

    def test_err_map_unchanged(self):
        error = DocError(code=ErrorCode.OPERATION_FAILED, message="Failed")
        result: Result[int] = Err(error)
        mapped = result.map(lambda x: x * 2)
        assert mapped.is_err()

    def test_ok_and_then(self):
        result: Result[int] = Ok(10)
        chained = result.and_then(lambda x: Ok(x + 5))
        assert chained.is_ok()
        assert chained.unwrap() == 15

    def test_err_and_then_short_circuits(self):
        error = DocError(code=ErrorCode.OPERATION_FAILED, message="Failed")
        result: Result[int] = Err(error)
        chained = result.and_then(lambda x: Ok(x + 5))
        assert chained.is_err()

    def test_collect_results_all_ok(self):
        results = [Ok(1), Ok(2), Ok(3)]
        collected = collect_results(results)
        assert collected.is_ok()
        assert collected.unwrap() == [1, 2, 3]

    def test_collect_results_with_error(self):
        error = DocError(code=ErrorCode.OPERATION_FAILED, message="Failed")
        results = [Ok(1), Err(error), Ok(3)]
        collected = collect_results(results)
        assert collected.is_err()

    def test_doc_error_with_context(self):
        error = DocError(code=ErrorCode.FILE_NOT_FOUND, message="Not found")
        enriched = error.with_context(path="/test/file.docx", line=42)
        assert enriched.details["path"] == "/test/file.docx"
        assert enriched.details["line"] == 42

    def test_doc_error_factory_methods(self):
        error = DocError.file_not_found("/path/to/file.docx")
        assert error.code == ErrorCode.FILE_NOT_FOUND
        assert "/path/to/file.docx" in error.message

    def test_try_operation_success(self):
        result = try_operation(lambda: 42, "Test operation")
        assert result.is_ok()
        assert result.unwrap() == 42

    def test_try_operation_failure(self):
        def failing():
            raise ValueError("Intentional error")

        result = try_operation(failing, "Test operation")
        assert result.is_err()
        assert "Intentional error" in result.unwrap_err().message


# =============================================================================
# Test DocumentEditor
# =============================================================================


class TestDocumentEditor:
    """Tests for DocumentEditor."""

    def test_open_nonexistent_file(self, temp_dir):
        path = temp_dir / "nonexistent.docx"
        editor = DocumentEditor(path)
        result = editor.open()
        assert result.is_err()
        assert result.unwrap_err().code == ErrorCode.FILE_NOT_FOUND

    def test_create_if_missing(self, temp_dir):
        path = temp_dir / "new.docx"
        editor = DocumentEditor(path, create_if_missing=True)
        result = editor.open()
        assert result.is_ok()
        assert editor.is_open
        assert editor.is_modified  # New doc is modified
        editor.close()

    def test_context_manager(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            assert editor.is_open
        assert not editor.is_open

    def test_add_paragraph(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            result = editor.add_paragraph("Hello, World!")
            assert result.is_ok()
            para_ref = result.unwrap()
            assert para_ref.text == "Hello, World!"
            assert para_ref.index >= 0

    def test_add_heading(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            result = editor.add_heading("Test Heading", level=1)
            assert result.is_ok()
            heading_ref = result.unwrap()
            assert heading_ref.text == "Test Heading"
            assert heading_ref.level == 1

    def test_add_table(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            result = editor.add_table(3, 4)
            assert result.is_ok()
            table_ref = result.unwrap()
            assert table_ref.rows == 3
            assert table_ref.cols == 4

    def test_fill_table(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            editor.add_table(3, 2)
            data = [["A", "B"], ["C", "D"], ["E", "F"]]
            result = editor.fill_table(0, data)
            assert result.is_ok()

    def test_get_stats(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            result = editor.get_stats()
            assert result.is_ok()
            stats = result.unwrap()
            assert stats["paragraph_count"] > 0
            assert stats["table_count"] == 1

    def test_get_structure(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            result = editor.get_structure()
            assert result.is_ok()
            structure = result.unwrap()
            # Should have headings (Title style isn't captured, only Heading styles)
            assert len(structure) > 0
            # First heading should be Chapter 1
            assert "Chapter 1" in structure[0]["text"]

    def test_find_paragraphs(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            result = editor.find_paragraphs(contains="chapter")
            assert result.is_ok()
            matches = result.unwrap()
            assert len(matches) > 0

    def test_replace_text(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            result = editor.replace_text("Python", "Rust")
            assert result.is_ok()
            count = result.unwrap()
            assert count >= 1

    def test_delete_paragraph(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            editor.add_paragraph("First")
            editor.add_paragraph("Second")
            editor.add_paragraph("Third")

            # Delete middle paragraph
            result = editor.delete_paragraph(1)
            assert result.is_ok()

            # Verify
            stats = editor.get_stats().unwrap()
            assert stats["paragraph_count"] == 2

    def test_invalid_position(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            result = editor.delete_paragraph(999)
            assert result.is_err()
            assert result.unwrap_err().code == ErrorCode.INVALID_POSITION


# =============================================================================
# Test Transactions
# =============================================================================


class TestTransactions:
    """Tests for transaction support."""

    def test_transaction_commit(self, empty_doc):
        with DocumentEditor(empty_doc, auto_save=False) as editor:
            with editor.transaction() as tx:
                editor.add_paragraph("Test content")
                tx.commit()

            # Changes should be kept
            stats = editor.get_stats().unwrap()
            assert stats["paragraph_count"] >= 1

    def test_transaction_rollback(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            initial_count = editor.get_stats().unwrap()["paragraph_count"]

            with editor.transaction() as tx:
                editor.add_paragraph("New content")
                tx.rollback()

            # Changes should be reverted
            final_count = editor.get_stats().unwrap()["paragraph_count"]
            assert final_count == initial_count

    def test_transaction_auto_rollback_on_exception(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            initial_count = editor.get_stats().unwrap()["paragraph_count"]

            try:
                with editor.transaction():
                    editor.add_paragraph("Will be rolled back")
                    raise ValueError("Intentional error")
            except ValueError:
                pass

            # Changes should be reverted
            final_count = editor.get_stats().unwrap()["paragraph_count"]
            assert final_count == initial_count


# =============================================================================
# Test High-Level Operations
# =============================================================================


class TestStructureReader:
    """Tests for StructureReader."""

    def test_get_table_of_contents(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            reader = StructureReader(editor)
            result = reader.get_table_of_contents()
            assert result.is_ok()
            toc = result.unwrap()
            assert len(toc) > 0

    def test_get_flat_toc(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            reader = StructureReader(editor)
            result = reader.get_flat_toc()
            assert result.is_ok()
            entries = result.unwrap()
            # Should have entries with level, text, index
            for entry in entries:
                assert "level" in entry
                assert "text" in entry
                assert "index" in entry

    def test_find_sections(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            reader = StructureReader(editor)
            result = reader.find_sections()
            assert result.is_ok()
            sections = result.unwrap()
            assert len(sections) > 0
            # Each section has heading and content range
            for section in sections:
                assert section.heading is not None
                assert section.content_start >= 0

    def test_find_section_by_heading(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            reader = StructureReader(editor)
            result = reader.find_section_by_heading("Chapter 1")
            assert result.is_ok()
            section = result.unwrap()
            assert section is not None
            assert "Chapter 1" in section.heading.text

    def test_get_chapter_index(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            reader = StructureReader(editor)
            result = reader.get_chapter_index()
            assert result.is_ok()
            chapters = result.unwrap()
            assert len(chapters) >= 2  # Chapter 1 and Chapter 2


class TestContentModifier:
    """Tests for ContentModifier."""

    def test_insert_after_heading(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            modifier = ContentModifier(editor)
            result = modifier.insert_after_heading(
                "Chapter 1",
                "Inserted paragraph after chapter 1"
            )
            assert result.is_ok()
            inserted = result.unwrap()
            assert len(inserted) == 1

    def test_insert_multiple_after_heading(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            modifier = ContentModifier(editor)
            result = modifier.insert_after_heading(
                "Chapter 1",
                ["First inserted", "Second inserted", "Third inserted"]
            )
            assert result.is_ok()
            inserted = result.unwrap()
            assert len(inserted) == 3

    def test_replace_in_section(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            modifier = ContentModifier(editor)
            result = modifier.replace_in_section(
                "Chapter 1",
                "basics",
                "fundamentals"
            )
            assert result.is_ok()

    def test_append_to_section(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            modifier = ContentModifier(editor)
            result = modifier.append_to_section(
                "Chapter 2",
                "Appended at the end of chapter 2"
            )
            assert result.is_ok()

    def test_insert_after_nonexistent_heading(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            modifier = ContentModifier(editor)
            result = modifier.insert_after_heading(
                "Nonexistent Chapter",
                "This should fail"
            )
            assert result.is_err()
            assert result.unwrap_err().code == ErrorCode.HEADING_NOT_FOUND


class TestTableBuilder:
    """Tests for TableBuilder."""

    def test_create_data_table(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            builder = TableBuilder(editor)
            data = [
                ["Name", "Age", "City"],
                ["Alice", "30", "NYC"],
                ["Bob", "25", "LA"],
            ]
            result = builder.create_data_table(data)
            assert result.is_ok()
            table_ref = result.unwrap()
            assert table_ref.rows == 3
            assert table_ref.cols == 3

    def test_create_key_value_table(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            builder = TableBuilder(editor)
            data = {
                "Version": "1.0.0",
                "Author": "Test",
                "Status": "Active",
            }
            result = builder.create_key_value_table(data)
            assert result.is_ok()

    def test_get_table_data(self, sample_doc):
        with DocumentEditor(sample_doc) as editor:
            builder = TableBuilder(editor)
            result = builder.get_table_data(0)
            assert result.is_ok()
            data = result.unwrap()
            assert len(data) == 3  # 3 rows
            assert data[0] == ["Name", "Value"]

    def test_add_row(self, sample_doc):
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            builder = TableBuilder(editor)
            result = builder.add_row(0, ["Gamma", "300"])
            assert result.is_ok()

            # Verify
            data = builder.get_table_data(0).unwrap()
            assert len(data) == 4  # Now 4 rows


# =============================================================================
# Test Planning System
# =============================================================================


class TestDocumentPlan:
    """Tests for DocumentPlan."""

    def test_create_plan(self):
        plan = DocumentPlan(
            name="Test Plan",
            description="A test plan"
        )
        assert plan.name == "Test Plan"
        assert plan.phase == PlanPhase.CREATED
        assert len(plan.steps) == 0

    def test_chainable_api(self):
        plan = (DocumentPlan("Chain Test")
            .add_heading("Title", level=1)
            .add_paragraph("Content")
            .add_table([["A", "B"], ["1", "2"]]))

        assert len(plan.steps) == 3

    def test_plan_to_dict(self):
        plan = DocumentPlan("Test").add_heading("Title", level=1)
        data = plan.to_dict()
        assert data["name"] == "Test"
        assert len(data["steps"]) == 1

    def test_plan_summary(self):
        plan = (DocumentPlan("Summary Test")
            .add_heading("Chapter", level=1)
            .add_paragraph("Content"))

        summary = plan.summary()
        assert "Summary Test" in summary
        assert "Chapter" in summary


class TestPlanExecutor:
    """Tests for PlanExecutor."""

    def test_execute_simple_plan(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            plan = (DocumentPlan("Simple Plan")
                .add_heading("Test Document", level=1)
                .add_paragraph("Hello, world!"))

            executor = PlanExecutor(editor)
            result = executor.execute(plan)

            assert result.is_ok()
            executed_plan = result.unwrap()
            assert executed_plan.phase == PlanPhase.COMPLETED

    def test_execute_with_callbacks(self, empty_doc):
        started = []
        completed = []

        with DocumentEditor(empty_doc) as editor:
            plan = (DocumentPlan("Callback Test")
                .add_heading("Title", level=1)
                .add_paragraph("Content"))

            executor = (PlanExecutor(editor)
                .on_step_start(lambda s: started.append(s.id))
                .on_step_complete(lambda s, r: completed.append(s.id)))

            executor.execute(plan)

            assert len(started) == 2
            assert len(completed) == 2

    def test_validate_plan(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            # Plan with missing required param
            plan = DocumentPlan("Invalid Plan")
            plan.add_step(ActionStep(
                action=ActionType.ADD_PARAGRAPH,
                params={},  # Missing 'text'
            ))

            executor = PlanExecutor(editor)
            result = executor.validate(plan)

            assert result.is_ok()
            warnings = result.unwrap()
            assert len(warnings) > 0
            assert "text" in warnings[0].lower()

    def test_dry_run(self, empty_doc):
        with DocumentEditor(empty_doc) as editor:
            plan = (DocumentPlan("Dry Run Test")
                .add_heading("Test", level=1))

            executor = PlanExecutor(editor)
            result = executor.execute(plan, dry_run=True)

            assert result.is_ok()
            executed_plan = result.unwrap()
            # Should not be completed (just validated)
            assert executed_plan.phase == PlanPhase.VALIDATING

    def test_plan_from_dict(self):
        data = {
            "name": "Dict Plan",
            "description": "Created from dict",
            "steps": [
                {
                    "action": "add_heading",
                    "params": {"text": "Title", "level": 1},
                },
                {
                    "action": "add_paragraph",
                    "params": {"text": "Content"},
                },
            ]
        }

        result = plan_from_dict(data)
        assert result.is_ok()
        plan = result.unwrap()
        assert plan.name == "Dict Plan"
        assert len(plan.steps) == 2

    def test_transaction_rollback_on_failure(self, sample_doc):
        """Test that failed plans rollback changes."""
        with DocumentEditor(sample_doc, auto_save=False) as editor:
            initial_count = editor.get_stats().unwrap()["paragraph_count"]

            # Create plan with invalid step
            plan = DocumentPlan("Failing Plan")
            plan.add_step(ActionStep(
                action=ActionType.ADD_PARAGRAPH,
                params={"text": "This will be added"},
            ))
            plan.add_step(ActionStep(
                action=ActionType.DELETE_PARAGRAPH,
                params={"index": 9999},  # Invalid index
            ))

            executor = PlanExecutor(editor)
            result = executor.execute(plan)

            # Plan should fail
            executed_plan = result.unwrap()
            assert executed_plan.phase == PlanPhase.FAILED

            # But changes should be rolled back
            final_count = editor.get_stats().unwrap()["paragraph_count"]
            assert final_count == initial_count


# =============================================================================
# Integration Test
# =============================================================================


class TestIntegration:
    """Integration tests combining multiple features."""

    def test_full_workflow(self, temp_dir):
        """Test creating a document from scratch with planning."""
        path = temp_dir / "full_workflow.docx"

        # Create document with plan
        with DocumentEditor(path, create_if_missing=True) as editor:
            plan = (DocumentPlan("Create Report")
                .add_heading("Quarterly Report", level=1)
                .add_paragraph("Executive summary goes here.")
                .add_heading("Financial Overview", level=2)
                .add_table([
                    ["Metric", "Q1", "Q2", "Q3"],
                    ["Revenue", "$1M", "$1.2M", "$1.5M"],
                    ["Costs", "$800K", "$900K", "$1M"],
                ])
                .add_heading("Conclusion", level=2)
                .add_paragraph("We exceeded expectations."))

            executor = PlanExecutor(editor)
            result = executor.execute(plan)
            assert result.is_ok()

            # Verify structure
            reader = StructureReader(editor)
            toc = reader.get_flat_toc().unwrap()
            assert len(toc) == 3  # 3 headings

            # Verify table
            builder = TableBuilder(editor)
            table_data = builder.get_table_data(0).unwrap()
            assert table_data[0][0] == "Metric"

        # Verify file was saved
        assert path.exists()

        # Re-open and verify persistence
        with DocumentEditor(path) as editor:
            stats = editor.get_stats().unwrap()
            assert stats["table_count"] == 1
            assert stats["paragraph_count"] >= 5
