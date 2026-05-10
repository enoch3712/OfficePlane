"""Tests for the python-docx renderer (Phase 10B)."""
from __future__ import annotations

import io
import json

from docx import Document

from officeplane.content_agent.renderers.blocks import (
    Block,
    BlocksDocument,
    Lesson,
    Module,
    parse_blocks_document,
)
from officeplane.content_agent.renderers.docx_blocks import render_docx


def test_renders_title_text_table_blocks():
    doc = BlocksDocument(
        title="Test Doc",
        modules=[
            Module(
                id="m1",
                title="Module One",
                lessons=[
                    Lesson(
                        id="l1",
                        title="Lesson A",
                        blocks=[
                            Block(type="title", content="Heading One", order=0),
                            Block(type="text", content="A paragraph of body text.", order=1),
                            Block(
                                type="table",
                                content=json.dumps({
                                    "headers": ["Col A", "Col B"],
                                    "rows": [["a1", "b1"], ["a2", "b2"]],
                                }),
                                order=2,
                            ),
                        ],
                    )
                ],
            )
        ],
    )
    blob = render_docx(doc)
    assert blob.startswith(b"PK")  # docx is a ZIP

    parsed = Document(io.BytesIO(blob))
    # level=0 → "Title" style; level≥1 → "Heading N" style
    titled = [p.text for p in parsed.paragraphs if p.style.name in ("Title",)]
    headings = [p.text for p in parsed.paragraphs if p.style.name.startswith("Heading")]
    assert "Test Doc" in titled
    assert "Module One" in headings
    assert "Lesson A" in headings
    assert "Heading One" in headings
    assert any("paragraph of body text" in p.text for p in parsed.paragraphs)
    assert len(parsed.tables) == 1
    assert parsed.tables[0].rows[0].cells[0].text == "Col A"


def test_parse_blocks_document_lenient():
    doc = parse_blocks_document({
        "title": "X",
        "modules": [{"id": "m", "lessons": [{"id": "l", "title": "L", "blocks": [
            {"type": "text", "content": "hi"},
            {"type": "unknown", "content": "should be dropped"},
        ]}]}],
    })
    assert doc.title == "X"
    blocks = doc.modules[0].lessons[0].blocks
    assert len(blocks) == 1
    assert blocks[0].type == "text"


def test_empty_blocks_document_renders_minimal_docx():
    blob = render_docx(BlocksDocument())
    assert blob.startswith(b"PK")
