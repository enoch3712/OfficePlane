"""Tests for the DOCX tree renderer (render_docx)."""

import io

from docx import Document as DocxDocument

from officeplane.content_agent.renderers.document import parse_document
from officeplane.content_agent.renderers.docx_render import render_docx


def test_render_docx_writes_headings_and_paragraphs():
    doc = parse_document(
        {
            "type": "document",
            "meta": {"title": "T"},
            "children": [
                {
                    "type": "section",
                    "level": 1,
                    "heading": "Intro",
                    "children": [
                        {"type": "paragraph", "text": "BP measurement matters."}
                    ],
                },
                {
                    "type": "section",
                    "level": 1,
                    "heading": "Method",
                    "children": [
                        {
                            "type": "section",
                            "level": 2,
                            "heading": "Cuff",
                            "children": [
                                {"type": "paragraph", "text": "Use correct size."},
                                {
                                    "type": "list",
                                    "ordered": True,
                                    "items": [
                                        {"type": "paragraph", "text": "Small"},
                                        {"type": "paragraph", "text": "Medium"},
                                    ],
                                },
                            ],
                        }
                    ],
                },
            ],
        }
    )
    blob = render_docx(doc)
    assert isinstance(blob, bytes) and len(blob) > 1000
    out = DocxDocument(io.BytesIO(blob))
    texts = [p.text for p in out.paragraphs]
    assert "Intro" in texts
    assert "Cuff" in texts
    assert any("Small" in t for t in texts)


def test_render_docx_emits_table():
    doc = parse_document(
        {
            "type": "document",
            "children": [
                {
                    "type": "table",
                    "headers": ["Systolic", "Diastolic"],
                    "rows": [["120", "80"], ["140", "90"]],
                }
            ],
        }
    )
    blob = render_docx(doc)
    out = DocxDocument(io.BytesIO(blob))
    assert len(out.tables) == 1
    assert out.tables[0].rows[0].cells[0].text == "Systolic"
    assert out.tables[0].rows[1].cells[1].text == "80"


def test_render_docx_handles_all_block_types_without_crashing():
    """Every block type renders without exception; Figure with no src is skipped."""
    doc = parse_document(
        {
            "type": "document",
            "meta": {"title": "All Blocks"},
            "children": [
                {"type": "heading", "level": 2, "text": "A heading"},
                {"type": "paragraph", "text": "A paragraph."},
                {
                    "type": "list",
                    "ordered": False,
                    "items": [
                        {"type": "paragraph", "text": "Bullet one"},
                        {"type": "paragraph", "text": "Bullet two"},
                    ],
                },
                {
                    "type": "table",
                    "headers": ["Name", "Value"],
                    "rows": [["alpha", "1"], ["beta", "2"]],
                },
                # Figure with no src — should be silently skipped
                {"type": "figure", "src": None, "caption": "Missing figure"},
                {"type": "code", "lang": "python", "text": "print('hello')"},
                {"type": "callout", "variant": "warning", "text": "Watch out!"},
                {"type": "quote", "text": "To be or not to be."},
                {"type": "divider"},
                {
                    "type": "section",
                    "level": 1,
                    "heading": "Nested section",
                    "children": [
                        {"type": "paragraph", "text": "Inside nested section."}
                    ],
                },
            ],
        }
    )
    blob = render_docx(doc)
    assert isinstance(blob, bytes) and len(blob) > 1000
    out = DocxDocument(io.BytesIO(blob))
    # At minimum the document has some paragraphs
    assert len(out.paragraphs) >= 1
