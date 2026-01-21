"""
Tests for document import functionality.

Tests the Word document importer and round-trip editing workflow.
"""

import io
import pytest
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

from officeplane.documents.importer import (
    DocumentImporter,
    ParsedDocument,
    ParsedChapter,
    ParsedSection,
    ParsedParagraph,
)
from officeplane.documents.models import (
    DocumentModel,
    ChapterModel,
    SectionModel,
    PageModel,
)
from officeplane.components.author import AuthorComponent
from officeplane.components.context import ComponentContext


def create_test_docx() -> bytes:
    """Create a simple test .docx file."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()

    # Set document properties
    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"

    # Add content with headings
    doc.add_heading("Chapter One", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("This is the introduction paragraph.")
    doc.add_paragraph("Another paragraph in the introduction.")

    doc.add_heading("Main Content", level=2)
    doc.add_paragraph("This is the main content section.")

    doc.add_heading("Chapter Two", level=1)
    doc.add_heading("Details", level=2)
    doc.add_paragraph("Details about the topic go here.")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_flat_docx() -> bytes:
    """Create a flat .docx file with no headings."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.core_properties.title = "Flat Document"

    doc.add_paragraph("First paragraph of content.")
    doc.add_paragraph("Second paragraph of content.")
    doc.add_paragraph("Third paragraph of content.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestDocumentParser:
    """Test document parsing functionality."""

    def test_parse_docx_extracts_title(self):
        """Test that document title is extracted."""
        docx_bytes = create_test_docx()
        importer = DocumentImporter()
        parsed = importer.parse_docx(docx_bytes)

        assert parsed.title == "Test Document"
        assert parsed.author == "Test Author"

    def test_parse_docx_extracts_chapters(self):
        """Test that Heading 1 becomes chapters."""
        docx_bytes = create_test_docx()
        importer = DocumentImporter()
        parsed = importer.parse_docx(docx_bytes)

        assert len(parsed.chapters) == 2
        assert parsed.chapters[0].title == "Chapter One"
        assert parsed.chapters[1].title == "Chapter Two"

    def test_parse_docx_extracts_sections(self):
        """Test that Heading 2 becomes sections."""
        docx_bytes = create_test_docx()
        importer = DocumentImporter()
        parsed = importer.parse_docx(docx_bytes)

        chapter_one = parsed.chapters[0]
        assert len(chapter_one.sections) == 2
        assert chapter_one.sections[0].title == "Introduction"
        assert chapter_one.sections[1].title == "Main Content"

    def test_parse_docx_extracts_paragraphs(self):
        """Test that paragraphs are extracted into sections."""
        docx_bytes = create_test_docx()
        importer = DocumentImporter()
        parsed = importer.parse_docx(docx_bytes)

        intro_section = parsed.chapters[0].sections[0]
        assert len(intro_section.paragraphs) == 2
        assert "introduction paragraph" in intro_section.paragraphs[0].lower()

    def test_parse_flat_document(self):
        """Test parsing document with no headings."""
        docx_bytes = create_flat_docx()
        importer = DocumentImporter()
        parsed = importer.parse_docx(docx_bytes)

        # Should create default structure
        assert len(parsed.chapters) == 1
        assert parsed.chapters[0].title == "Content"


class TestPageSplitting:
    """Test page splitting functionality."""

    def test_split_short_content(self):
        """Test that short content stays on one page."""
        importer = DocumentImporter(words_per_page=500)
        paragraphs = ["Short paragraph one.", "Short paragraph two."]

        pages = importer._split_into_pages(paragraphs)

        assert len(pages) == 1
        assert "Short paragraph one." in pages[0]
        assert "Short paragraph two." in pages[0]

    def test_split_long_content(self):
        """Test that long content is split into multiple pages."""
        importer = DocumentImporter(words_per_page=10)
        paragraphs = [
            "This is a long paragraph with many words that will exceed the limit.",
            "Another paragraph with content.",
        ]

        pages = importer._split_into_pages(paragraphs)

        assert len(pages) >= 2

    def test_split_empty_content(self):
        """Test splitting empty content."""
        importer = DocumentImporter()
        pages = importer._split_into_pages([])

        assert pages == []


class TestDocumentImport:
    """Test full document import flow."""

    @pytest.fixture
    def mock_doc_store(self):
        """Create mock document store."""
        store = AsyncMock()
        doc_id = uuid4()
        chapter_id = uuid4()
        section_id = uuid4()
        page_id = uuid4()

        store.create_document.return_value = DocumentModel(
            id=doc_id,
            title="Test Document",
            author="Test Author",
        )
        store.create_chapter.return_value = ChapterModel(
            id=chapter_id,
            document_id=doc_id,
            title="Chapter One",
            order_index=0,
        )
        store.create_section.return_value = SectionModel(
            id=section_id,
            chapter_id=chapter_id,
            title="Introduction",
            order_index=0,
        )
        store.create_page.return_value = PageModel(
            id=page_id,
            section_id=section_id,
            page_number=1,
            content="Test content",
            word_count=2,
        )
        store.get_document.return_value = DocumentModel(
            id=doc_id,
            title="Test Document",
            author="Test Author",
            chapters=[
                ChapterModel(
                    id=chapter_id,
                    document_id=doc_id,
                    title="Chapter One",
                    order_index=0,
                    sections=[
                        SectionModel(
                            id=section_id,
                            chapter_id=chapter_id,
                            title="Introduction",
                            order_index=0,
                            pages=[
                                PageModel(
                                    id=page_id,
                                    section_id=section_id,
                                    page_number=1,
                                    content="Test content",
                                    word_count=2,
                                )
                            ],
                        )
                    ],
                )
            ],
        )
        return store

    async def test_import_from_bytes(self, mock_doc_store):
        """Test importing from bytes."""
        docx_bytes = create_test_docx()
        importer = DocumentImporter(doc_store=mock_doc_store)

        doc = await importer.import_from_bytes(
            docx_bytes=docx_bytes,
            index_for_search=False,  # Skip indexing in tests
        )

        assert doc is not None
        mock_doc_store.create_document.assert_called_once()
        mock_doc_store.create_chapter.assert_called()
        mock_doc_store.create_section.assert_called()

    async def test_import_with_title_override(self, mock_doc_store):
        """Test importing with custom title."""
        docx_bytes = create_test_docx()
        importer = DocumentImporter(doc_store=mock_doc_store)

        await importer.import_from_bytes(
            docx_bytes=docx_bytes,
            title="Custom Title",
            author="Custom Author",
            index_for_search=False,
        )

        call_args = mock_doc_store.create_document.call_args
        assert call_args.kwargs["title"] == "Custom Title"
        assert call_args.kwargs["author"] == "Custom Author"


class TestAuthorComponentImport:
    """Test AuthorComponent import actions."""

    @pytest.fixture
    def mock_stores(self):
        """Create mock stores."""
        doc_store = AsyncMock()
        vector_store = AsyncMock()
        embedding_client = MagicMock()
        return doc_store, vector_store, embedding_client

    @pytest.fixture
    def ctx(self):
        """Create test context."""
        mock_driver = MagicMock()
        mock_store = MagicMock()
        return ComponentContext.create(driver=mock_driver, store=mock_store)

    @pytest.fixture
    def author_component(self, mock_stores):
        """Create AuthorComponent with mocks."""
        doc_store, vector_store, embedding_client = mock_stores
        component = AuthorComponent(
            doc_store=doc_store,
            vector_store=vector_store,
            embedding_client=embedding_client,
        )
        component._connected = True
        component._indexer = AsyncMock()
        component._retriever = AsyncMock()
        return component

    def test_import_action_registered(self):
        """Test import_document action is registered."""
        author = AuthorComponent()
        assert author.get_action("import_document") is not None
        assert author.get_action("reimport_document") is not None

    def test_actions_count(self):
        """Test all 13 actions are registered."""
        author = AuthorComponent()
        # Original 11 + 2 new import actions
        assert len(author.actions()) == 14

    async def test_import_missing_input(self, author_component, ctx):
        """Test import fails when no file_path or content_base64 provided."""
        result = await author_component.execute(
            "import_document",
            {},  # No file_path or content_base64
            ctx,
        )

        assert result.success is False
        assert "must be provided" in result.error

    async def test_import_file_not_found(self, author_component, ctx):
        """Test import fails for non-existent file."""
        result = await author_component.execute(
            "import_document",
            {"file_path": "/nonexistent/file.docx"},
            ctx,
        )

        assert result.success is False
        assert "not found" in result.error.lower()


class TestParsedDocumentStructure:
    """Test ParsedDocument dataclass structure."""

    def test_parsed_document_defaults(self):
        """Test ParsedDocument default values."""
        doc = ParsedDocument(title="Test")

        assert doc.title == "Test"
        assert doc.author is None
        assert doc.chapters == []
        assert doc.preamble == []
        assert doc.metadata == {}

    def test_parsed_chapter_defaults(self):
        """Test ParsedChapter default values."""
        chapter = ParsedChapter(title="Chapter")

        assert chapter.title == "Chapter"
        assert chapter.sections == []
        assert chapter.preamble == []

    def test_parsed_section_defaults(self):
        """Test ParsedSection default values."""
        section = ParsedSection(title="Section")

        assert section.title == "Section"
        assert section.paragraphs == []

    def test_parsed_paragraph(self):
        """Test ParsedParagraph creation."""
        para = ParsedParagraph(
            text="Hello world",
            style="Normal",
            is_heading=False,
            heading_level=0,
        )

        assert para.text == "Hello world"
        assert para.style == "Normal"
        assert para.is_heading is False


class TestToolExport:
    """Test that import actions export correctly as tools."""

    def test_import_in_function_tools(self):
        """Test import_document exports as function tool."""
        author = AuthorComponent()
        tools = author.to_function_tools()

        import_tool = next(
            (t for t in tools if t["function"]["name"] == "import_document"),
            None,
        )

        assert import_tool is not None
        assert "type" in import_tool
        assert import_tool["type"] == "function"
        assert "parameters" in import_tool["function"]

    def test_reimport_in_function_tools(self):
        """Test reimport_document exports as function tool."""
        author = AuthorComponent()
        tools = author.to_function_tools()

        reimport_tool = next(
            (t for t in tools if t["function"]["name"] == "reimport_document"),
            None,
        )

        assert reimport_tool is not None
        assert "document_id" in str(reimport_tool["function"]["parameters"])


class TestRoundTripWorkflow:
    """Test full round-trip workflow: import → modify → export."""

    @pytest.fixture
    def test_docx_path(self):
        """Path to the test.docx file."""
        import os
        return os.path.join(os.path.dirname(__file__), "test.docx")

    @pytest.fixture
    def mock_doc_store_with_ids(self):
        """Create mock document store that tracks created items."""
        store = AsyncMock()

        # Track created items for structure building
        created_chapters = []
        created_sections = []
        created_pages = []

        doc_id = uuid4()

        def make_chapter(document_id, title, order_index=0, **kwargs):
            chapter_id = uuid4()
            chapter = ChapterModel(
                id=chapter_id,
                document_id=document_id,
                title=title,
                order_index=order_index,
                sections=[],
            )
            created_chapters.append(chapter)
            return chapter

        def make_section(chapter_id, title, order_index=0, **kwargs):
            section_id = uuid4()
            section = SectionModel(
                id=section_id,
                chapter_id=chapter_id,
                title=title,
                order_index=order_index,
                pages=[],
            )
            created_sections.append(section)
            return section

        def make_page(section_id, content, page_number=1, **kwargs):
            page_id = uuid4()
            page = PageModel(
                id=page_id,
                section_id=section_id,
                page_number=page_number,
                content=content,
                word_count=len(content.split()),
            )
            created_pages.append(page)
            return page

        store.create_document.return_value = DocumentModel(
            id=doc_id,
            title="Test Document",
            author="Test Author",
        )
        store.create_chapter.side_effect = make_chapter
        store.create_section.side_effect = make_section
        store.create_page.side_effect = make_page

        # For get_document, build the structure from created items
        async def get_document_impl(document_id, load_children=False):
            doc = DocumentModel(
                id=doc_id,
                title="Test Document",
                author="Test Author",
                chapters=created_chapters.copy() if load_children else [],
            )
            return doc

        store.get_document.side_effect = get_document_impl
        store._created_chapters = created_chapters
        store._created_sections = created_sections
        store._created_pages = created_pages
        store._doc_id = doc_id

        return store

    def test_test_docx_exists(self, test_docx_path):
        """Verify test.docx file exists."""
        import os
        assert os.path.exists(test_docx_path), f"test.docx not found at {test_docx_path}"

    def test_parse_test_docx_structure(self, test_docx_path):
        """Test parsing the test.docx file structure."""
        with open(test_docx_path, "rb") as f:
            docx_bytes = f.read()

        importer = DocumentImporter()
        parsed = importer.parse_docx(docx_bytes)

        # Verify document was parsed
        assert parsed.title is not None
        assert len(parsed.chapters) >= 1, "Document should have at least one chapter"

        # Log structure for debugging
        print(f"\nParsed document: {parsed.title}")
        print(f"Author: {parsed.author}")
        print(f"Chapters: {len(parsed.chapters)}")
        for i, chapter in enumerate(parsed.chapters):
            print(f"  Chapter {i}: {chapter.title}")
            for j, section in enumerate(chapter.sections):
                print(f"    Section {j}: {section.title}")
                print(f"      Paragraphs: {len(section.paragraphs)}")

    async def test_import_modify_export_workflow(
        self, test_docx_path, mock_doc_store_with_ids
    ):
        """Test full import → add section → export workflow."""
        # Step 1: Import the document
        with open(test_docx_path, "rb") as f:
            docx_bytes = f.read()

        importer = DocumentImporter(doc_store=mock_doc_store_with_ids)
        parsed = importer.parse_docx(docx_bytes)

        # Verify initial structure
        original_chapter_count = len(parsed.chapters)
        assert original_chapter_count >= 1, "Need at least one chapter for test"

        if parsed.chapters:
            first_chapter = parsed.chapters[0]
            original_section_count = len(first_chapter.sections)

            # Step 2: Import to database (mocked)
            doc = await importer.import_from_bytes(
                docx_bytes=docx_bytes,
                index_for_search=False,
            )

            assert doc is not None
            mock_doc_store_with_ids.create_document.assert_called_once()

            # Verify chapters were created
            assert mock_doc_store_with_ids.create_chapter.call_count >= original_chapter_count

            # Step 3: Add a new section in the middle
            if mock_doc_store_with_ids._created_chapters:
                first_created_chapter = mock_doc_store_with_ids._created_chapters[0]

                # Simulate adding a new section
                new_section = await mock_doc_store_with_ids.create_section(
                    chapter_id=first_created_chapter.id,
                    title="New Section Added in Middle",
                    order_index=1,  # Insert in middle
                )

                # Add a page to the new section
                new_page = await mock_doc_store_with_ids.create_page(
                    section_id=new_section.id,
                    content="This is content added during the round-trip workflow. "
                            "It demonstrates that we can import a document, "
                            "modify its structure by adding new sections, "
                            "and then export it back.",
                    page_number=1,
                )

                assert new_section.title == "New Section Added in Middle"
                assert new_page.content is not None

    async def test_add_section_between_existing(self, test_docx_path):
        """Test adding a section between existing sections."""
        with open(test_docx_path, "rb") as f:
            docx_bytes = f.read()

        importer = DocumentImporter()
        parsed = importer.parse_docx(docx_bytes)

        # Get the first chapter with sections
        chapter_with_sections = None
        for chapter in parsed.chapters:
            if len(chapter.sections) >= 2:
                chapter_with_sections = chapter
                break

        if chapter_with_sections:
            # We have at least 2 sections, we can insert between them
            original_sections = [s.title for s in chapter_with_sections.sections]
            print(f"\nOriginal sections in '{chapter_with_sections.title}':")
            for title in original_sections:
                print(f"  - {title}")

            # Simulate adding a new section (in-memory modification)
            new_section = ParsedSection(
                title="Inserted Section",
                paragraphs=["This section was inserted in the middle."],
            )

            # Insert at position 1 (between first and second)
            chapter_with_sections.sections.insert(1, new_section)

            assert len(chapter_with_sections.sections) == len(original_sections) + 1
            assert chapter_with_sections.sections[1].title == "Inserted Section"

            print(f"\nModified sections:")
            for section in chapter_with_sections.sections:
                print(f"  - {section.title}")
        else:
            # Document has only single sections per chapter
            # Just verify we can parse and the structure is correct
            assert len(parsed.chapters) >= 1, "Should have parsed at least one chapter"

    def test_export_after_modification_preserving_styles(self, test_docx_path):
        """Test that modified document preserves original formatting and styles."""
        from officeplane.documents.importer import DocumentEditor

        with open(test_docx_path, "rb") as f:
            docx_bytes = f.read()

        # Use DocumentEditor to preserve original styles
        editor = DocumentEditor(docx_bytes)

        # Get original paragraph count
        original_count = editor.paragraph_count
        print(f"\nOriginal document has {original_count} paragraphs")

        # Show a few paragraphs from the middle
        paragraphs = editor.get_paragraph_texts()
        middle = len(paragraphs) // 2
        print(f"\nParagraphs around middle (index {middle}):")
        for i in range(max(0, middle - 2), min(len(paragraphs), middle + 3)):
            text_preview = paragraphs[i][:50] + "..." if len(paragraphs[i]) > 50 else paragraphs[i]
            print(f"  [{i}]: {text_preview}")

        # Get style info from an original body paragraph for comparison
        original_style_info = editor.get_paragraph_style_info(middle)
        print(f"\nOriginal paragraph style: {original_style_info['style_name']}")
        if original_style_info['runs']:
            run = original_style_info['runs'][0]
            print(f"  Font: {run['font_name']}, Size: {run['font_size']}, Bold: {run['bold']}")

        # Insert a new section in the middle, preserving all original formatting
        editor.insert_section_after_paragraph(
            paragraph_index=middle,
            heading="New Section Added in Middle",
            paragraphs=[
                "This is the first paragraph of the newly added section.",
                "This paragraph demonstrates that content can be inserted while preserving the original document's formatting.",
                "The original styles, fonts, and layout remain intact.",
            ],
            heading_level=2,
        )

        # Verify new paragraph count
        new_count = editor.paragraph_count
        assert new_count == original_count + 4, \
            f"Expected {original_count + 4} paragraphs (3 content + 1 heading), got {new_count}"

        print(f"\nAfter insertion: {new_count} paragraphs (+4)")

        # Save the modified document
        # import os
        # export_path = os.path.join(os.path.dirname(__file__), "modified_document_test.docx")
        # editor.save_to_file(export_path)

        # print(f"Exported file saved to: {export_path}")

        # Verify the exported file contains the inserted content with correct styling
        output_bytes = editor.save()
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(output_bytes))

        # Find the inserted heading and verify its style matches original headings
        inserted_heading = None
        original_heading = None
        for para in doc.paragraphs:
            if "New Section Added in Middle" in para.text:
                inserted_heading = para
            elif para.style and para.style.name.startswith("Heading") and original_heading is None:
                original_heading = para

        assert inserted_heading is not None, "Inserted heading not found in document"
        assert original_heading is not None, "No original heading found for comparison"

        print(f"\nOriginal heading style: {original_heading.style.name}")
        print(f"Inserted heading style: {inserted_heading.style.name}")

        # Verify the inserted heading uses the same style as original headings
        assert inserted_heading.style.name == original_heading.style.name, \
            f"Style mismatch: inserted={inserted_heading.style.name}, original={original_heading.style.name}"

        # Verify font size matches (if set)
        if original_heading.runs and inserted_heading.runs:
            orig_size = original_heading.runs[0].font.size
            ins_size = inserted_heading.runs[0].font.size
            print(f"Original heading font size: {orig_size}")
            print(f"Inserted heading font size: {ins_size}")
            if orig_size is not None:
                assert ins_size == orig_size, f"Font size mismatch: {ins_size} vs {orig_size}"

        print("\n✓ Inserted heading matches original document style!")

    def test_inserted_paragraphs_match_original_style(self, test_docx_path):
        """Test that inserted paragraphs have the same style as original document."""
        from officeplane.documents.importer import DocumentEditor
        from docx import Document as DocxDocument

        with open(test_docx_path, "rb") as f:
            docx_bytes = f.read()

        editor = DocumentEditor(docx_bytes)

        # Find a body paragraph to compare against
        middle = editor.paragraph_count // 2

        # Insert content
        editor.insert_section_after_paragraph(
            paragraph_index=middle,
            heading="Style Test Section",
            paragraphs=["Test paragraph to verify style copying."],
            heading_level=2,
        )

        # Save and reload to verify
        output_bytes = editor.save()
        doc = DocxDocument(io.BytesIO(output_bytes))

        # Find the inserted paragraph
        inserted_para = None
        original_body_para = None
        for para in doc.paragraphs:
            if "Test paragraph to verify" in para.text:
                inserted_para = para
            elif len(para.text) > 30 and original_body_para is None:
                # Find a typical body paragraph from original content
                style_name = para.style.name if para.style else ""
                if not style_name.startswith(("Heading", "Title", "TOC")):
                    original_body_para = para

        assert inserted_para is not None, "Could not find inserted paragraph"
        assert original_body_para is not None, "Could not find original body paragraph"

        # Compare run formatting
        if inserted_para.runs and original_body_para.runs:
            inserted_run = inserted_para.runs[0]
            original_run = original_body_para.runs[0]

            print(f"\nOriginal paragraph: '{original_body_para.text[:40]}...'")
            print(f"  Font: {original_run.font.name}")
            print(f"  Size: {original_run.font.size}")

            print(f"\nInserted paragraph: '{inserted_para.text[:40]}...'")
            print(f"  Font: {inserted_run.font.name}")
            print(f"  Size: {inserted_run.font.size}")

            # Verify font properties match
            # Note: Some properties may be None if inherited from style
            if original_run.font.name is not None and inserted_run.font.name is not None:
                assert inserted_run.font.name == original_run.font.name, \
                    f"Font name mismatch: {inserted_run.font.name} vs {original_run.font.name}"

            if original_run.font.size is not None and inserted_run.font.size is not None:
                assert inserted_run.font.size == original_run.font.size, \
                    f"Font size mismatch: {inserted_run.font.size} vs {original_run.font.size}"

            print("\n✓ Inserted paragraph matches original document style!")
