"""DOCX renderer that walks the agnostic Document tree and produces .docx bytes.

Uses python-docx to emit a well-structured Word document.  The renderer is
intentionally stateless — each call to :func:`render_docx` opens a fresh
``python_docx.Document`` and serialises it to a ``BytesIO`` buffer.
"""

from __future__ import annotations

import logging
from io import BytesIO
from os.path import isfile
from typing import Union

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from officeplane.content_agent.renderers.document import (
    Callout,
    Code,
    Divider,
    Document,
    Figure,
    Heading,
    List,
    Paragraph,
    Quote,
    Section,
    Table,
)

logger = logging.getLogger(__name__)

# Union of all node types the walker encounters
_Node = Union[Section, Heading, Paragraph, List, Table, Figure, Code, Callout, Quote, Divider]


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _safe_style(doc_obj: DocxDocument, name: str) -> str | None:
    """Return *name* if it exists in *doc_obj*'s style catalogue, else None."""
    try:
        _ = doc_obj.styles[name]
        return name
    except KeyError:
        return None


def _render_node(doc_obj: DocxDocument, node: _Node) -> None:
    """Dispatch a single node to its specialised render function."""
    if isinstance(node, Section):
        _render_section(doc_obj, node)
    elif isinstance(node, Heading):
        doc_obj.add_heading(node.text, level=min(node.level, 9))
    elif isinstance(node, Paragraph):
        doc_obj.add_paragraph(node.text)
    elif isinstance(node, List):
        _render_list(doc_obj, node)
    elif isinstance(node, Table):
        _render_table(doc_obj, node)
    elif isinstance(node, Figure):
        _render_figure(doc_obj, node)
    elif isinstance(node, Code):
        _render_code(doc_obj, node)
    elif isinstance(node, Callout):
        _render_callout(doc_obj, node)
    elif isinstance(node, Quote):
        _render_quote(doc_obj, node)
    elif isinstance(node, Divider):
        doc_obj.add_page_break()
    else:
        logger.warning("docx_render: unknown node type %s — skipped", type(node).__name__)


def _render_section(doc_obj: DocxDocument, section: Section) -> None:
    """Emit a heading for the section then recurse into its children."""
    if section.heading:
        doc_obj.add_heading(section.heading, level=min(section.level, 9))
    for child in section.children:
        _render_node(doc_obj, child)


def _render_list(doc_obj: DocxDocument, lst: List) -> None:
    style = "List Number" if lst.ordered else "List Bullet"
    for item in lst.items:
        doc_obj.add_paragraph(item.text, style=style)


def _render_table(doc_obj: DocxDocument, table: Table) -> None:
    n_cols = len(table.headers)
    n_rows = len(table.rows) + 1  # +1 for header row

    tbl_style = _safe_style(doc_obj, "Table Grid") or "Normal Table"
    try:
        tbl = doc_obj.add_table(rows=n_rows, cols=n_cols, style=tbl_style)
    except KeyError:
        tbl = doc_obj.add_table(rows=n_rows, cols=n_cols)

    # Header row — bold
    hdr_cells = tbl.rows[0].cells
    for i, header_text in enumerate(table.headers):
        hdr_cells[i].text = header_text
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for row_idx, row_data in enumerate(table.rows):
        data_cells = tbl.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < n_cols:
                data_cells[col_idx].text = cell_text


def _render_figure(doc_obj: DocxDocument, figure: Figure) -> None:
    """Embed a figure if src points to an existing file; silently skip otherwise."""
    src = figure.src
    if not src or not isfile(src):
        # Missing or absent — skip gracefully (Task 8 image_embed will fill src)
        logger.debug("docx_render: figure %s skipped — src absent or file missing", figure.id)
        return

    try:
        doc_obj.add_picture(src, width=Inches(5))
    except Exception:
        logger.warning(
            "docx_render: could not embed figure %s from %s", figure.id, src, exc_info=True
        )
        return

    if figure.caption:
        p = doc_obj.add_paragraph()
        r = p.add_run(figure.caption)
        r.italic = True


def _render_code(doc_obj: DocxDocument, code: Code) -> None:
    """Render a code block as a paragraph with Courier New font.

    Prefixes a ``[lang]`` tag when ``code.lang`` is set.
    """
    text = f"[{code.lang}]\n{code.text}" if code.lang else code.text
    para = doc_obj.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Courier New"
    # Ensure the font is applied in the XML (python-docx quirk for certain fonts)
    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Courier New")
    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Courier New")


def _render_callout(doc_obj: DocxDocument, callout: Callout) -> None:
    """Render a callout as a prefixed paragraph with Intense Quote style if available."""
    text = f"[{callout.variant.upper()}] {callout.text}"
    style = _safe_style(doc_obj, "Intense Quote")
    try:
        if style:
            doc_obj.add_paragraph(text, style=style)
        else:
            doc_obj.add_paragraph(text)
    except Exception:
        doc_obj.add_paragraph(text)


def _render_quote(doc_obj: DocxDocument, quote: Quote) -> None:
    """Render a block quote using Intense Quote style with plain fallback."""
    try:
        doc_obj.add_paragraph(quote.text, style="Intense Quote")
    except KeyError:
        doc_obj.add_paragraph(quote.text)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def render_docx(doc: Document) -> bytes:
    """Render the agnostic Document tree to .docx bytes via python-docx.

    Opens a fresh :class:`~docx.Document`, writes the document title (if
    present in ``doc.meta.title``) as a level-0 heading, then walks
    ``doc.children`` depth-first dispatching each node to the appropriate
    render helper.  The result is serialised to a :class:`~io.BytesIO` buffer
    and returned as raw bytes.
    """
    doc_obj = DocxDocument()

    # Document title
    if doc.meta.title and doc.meta.title != "Untitled":
        doc_obj.add_heading(doc.meta.title, level=0)

    # Walk children
    for node in doc.children:
        _render_node(doc_obj, node)

    buf = BytesIO()
    doc_obj.save(buf)
    return buf.getvalue()
