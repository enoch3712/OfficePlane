"""Render a BlocksDocument to a .docx via python-docx."""
from __future__ import annotations

import io
import json
import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from officeplane.content_agent.renderers.blocks import (
    Block,
    BlocksDocument,
    Lesson,
    Module,
)

log = logging.getLogger("officeplane.content_agent.renderers.docx_blocks")


def render_docx(doc: BlocksDocument, *, image_root: Path | None = None) -> bytes:
    """Render the blocks doc to .docx bytes."""
    document = Document()

    # Document title
    if doc.title and doc.title != "Untitled":
        document.add_heading(doc.title, level=0)

    for module in doc.modules:
        if module.title:
            document.add_heading(module.title, level=1)
        for lesson in module.lessons:
            if lesson.title:
                document.add_heading(lesson.title, level=2)
            for block in sorted(lesson.blocks, key=lambda b: b.order):
                _render_block(document, block, image_root=image_root)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _render_block(document, block: Block, *, image_root: Path | None) -> None:
    if block.type == "title" and block.content:
        document.add_heading(block.content, level=3)
        _emit_source_refs(document, block)
        return
    if block.type == "text" and block.content:
        document.add_paragraph(block.content)
        _emit_source_refs(document, block)
        return
    if block.type == "table" and block.content:
        try:
            spec = json.loads(block.content)
        except json.JSONDecodeError:
            log.warning("table block content is not JSON; skipping")
            return
        _render_table(document, spec)
        _emit_source_refs(document, block)
        return
    if block.type == "image" and block.object_key and image_root is not None:
        path = image_root / block.object_key
        if path.exists():
            document.add_picture(str(path))
        elif block.alt:
            document.add_paragraph(f"[Image: {block.alt}]")
        _emit_source_refs(document, block)
        return


def _render_table(document, spec: dict) -> None:
    headers = list(spec.get("headers") or [])
    rows = [list(r) for r in (spec.get("rows") or [])]
    if not headers and not rows:
        return
    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return
    table = document.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Light Grid"
    if headers:
        head = table.rows[0]
        for i, h in enumerate(headers):
            head.cells[i].text = str(h)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)


def _emit_source_refs(document, block: Block) -> None:
    if not block.source_references:
        return
    refs_text = "; ".join(
        f"{r.document_title or 'Document'}"
        + (f" · {r.chapter_title}" if r.chapter_title else "")
        + (f" · {r.section_title}" if r.section_title else "")
        + (f" (pp. {','.join(str(p) for p in r.page_numbers)})" if r.page_numbers else "")
        for r in block.source_references
    )
    p = document.add_paragraph()
    run = p.add_run(f"Source: {refs_text}")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
