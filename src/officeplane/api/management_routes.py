"""
FastAPI routes for OfficePlane Management System
Instances, Tasks, History, Metrics, and WebSocket
"""
from fastapi import APIRouter, HTTPException, WebSocket, WebSocketDisconnect, UploadFile, File, Body, Response
from typing import Optional, List, Dict, Any
from datetime import datetime
import asyncio
import json
from pydantic import BaseModel

from prisma.enums import InstanceState, TaskState, TaskPriority, EventType
from ..management.db import get_db
from ..management.instance_manager import instance_manager
from ..management.task_queue import task_queue
from ..documents.models import DocumentModel
from ..documents.store import DocumentStore
from ..ingestion.config import IngestionConfig
from ..ingestion.format_detector import DocumentFormat, detect_format
# VisionIngestionService and LibreOfficeDriver are imported lazily to avoid requiring PIL at startup
from ..components.planning.generator import PlanGenerator, MockPlanLLM, GeminiPlanAdapter
from ..components.planning.models import GeneratePlanInput, PlanSummary, ActionPlan, ActionNode
from ..components.planning.display import PlanDisplayer
from ..components.runner import OpenAIAdapter
from ..config import config

router = APIRouter(prefix="/api")

# WebSocket connections for real-time updates
websocket_connections: set[WebSocket] = set()


# ============================================================
# REQUEST/RESPONSE MODELS
# ============================================================


class CreateInstanceRequest(BaseModel):
    documentId: Optional[str] = None
    driverType: Optional[str] = None
    filePath: Optional[str] = None


class PlanDocumentRequest(BaseModel):
    prompt: str
    max_chapters: int = 20
    max_sections_per_chapter: int = 10
    max_pages_per_section: int = 5
    include_content_outlines: bool = True


async def broadcast_event(event_type: str, data: dict):
    """Broadcast event to all connected WebSocket clients"""
    message = {"type": event_type, "data": data, "timestamp": datetime.utcnow().isoformat()}

    disconnected = set()
    for ws in websocket_connections:
        try:
            await ws.send_json(message)
        except Exception:
            disconnected.add(ws)

    # Clean up disconnected clients
    websocket_connections.difference_update(disconnected)


def _format_outline_for_prompt(outline) -> str:
    lines: List[str] = []
    lines.append(f"Document: {outline.title} ({outline.id})")
    if outline.author:
        lines.append(f"Author: {outline.author}")
    lines.append(
        f"Chapters: {outline.chapter_count}, Sections: {outline.section_count}, Pages: {outline.page_count}"
    )
    for chapter in outline.chapters:
        lines.append(f"- Chapter {chapter.order_index + 1}: {chapter.title} ({chapter.id})")
        for section in chapter.sections:
            lines.append(
                f"  - Section {section.order_index + 1}: {section.title} ({section.id})"
            )
            for page in section.pages:
                lines.append(f"    - Page {page.page_number} ({page.id})")
    return "\n".join(lines)


def _serialize_document(document: DocumentModel) -> dict:
    chapters = document.chapters or []
    return {
        "id": str(document.id),
        "title": document.title,
        "author": document.author,
        "chapters": [
            {
                "id": str(ch.id),
                "title": ch.title,
                "order_index": ch.order_index,
                "sections": [
                    {
                        "id": str(sec.id),
                        "title": sec.title,
                        "order_index": sec.order_index,
                        "page_count": len(sec.pages) if sec.pages else 0,
                    }
                    for sec in (ch.sections or [])
                ],
            }
            for ch in chapters
        ],
        "total_chapters": len(chapters),
        "total_pages": sum(
            len(sec.pages or [])
            for ch in chapters
            for sec in (ch.sections or [])
        ),
    }


def _normalize_plan_for_existing_document(
    plan: ActionPlan,
    document_id: str,
    title: str,
    existing_chapter_titles: Optional[set[str]] = None,
) -> ActionPlan:
    roots: List[ActionNode] = []

    def normalize_node(node: ActionNode, parent_id: Optional[str]) -> None:
        node.parent_id = parent_id
        for idx, child in enumerate(node.children):
            if child.order_index == 0:
                child.order_index = idx
            normalize_node(child, node.id)

        if node.action_name == "add_chapter":
            node.inputs["document_id"] = document_id
        elif node.action_name == "write_page":
            if "content" not in node.inputs and "content_outline" in node.inputs:
                node.inputs["content"] = node.inputs.pop("content_outline")
            node.inputs.pop("title", None)
            node.inputs.setdefault("content", "")

    existing_titles = {t.lower() for t in (existing_chapter_titles or set())}

    for root in plan.roots:
        if root.action_name == "create_document":
            for child in root.children:
                if (
                    child.action_name == "add_chapter"
                    and child.inputs.get("title", "").lower() in existing_titles
                ):
                    continue
                child.parent_id = None
                normalize_node(child, None)
                roots.append(child)
        else:
            root.parent_id = None
            if (
                root.action_name == "add_chapter"
                and root.inputs.get("title", "").lower() in existing_titles
            ):
                continue
            normalize_node(root, None)
            roots.append(root)

    return ActionPlan(
        title=f"Plan for {title}",
        original_prompt=plan.original_prompt,
        roots=roots,
    )


def _build_plan_generator() -> PlanGenerator:
    """Build a PlanGenerator with the best available LLM.

    Tries in order: Gemini (GOOGLE_API_KEY), OpenAI (OPENAI_API_KEY), MockPlanLLM (fallback).
    """
    import os
    import logging

    # Try Gemini first (used for vision ingestion, likely already configured)
    google_key = os.getenv("GOOGLE_API_KEY")
    if google_key:
        try:
            model = os.getenv("OFFICEPLANE_PLAN_MODEL", "gemini-2.0-flash")
            adapter = GeminiPlanAdapter(api_key=google_key, model=model)
            logging.info(f"Using Gemini for plan generation (model: {model})")
            return PlanGenerator(llm=adapter)
        except Exception as e:
            logging.warning(f"Failed to initialize Gemini adapter: {e}")

    # Try OpenAI as fallback
    openai_key = os.getenv("OPENAI_API_KEY")
    if openai_key:
        try:
            from openai import AsyncOpenAI

            client = AsyncOpenAI(api_key=openai_key)
            model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
            logging.info(f"Using OpenAI for plan generation (model: {model})")
            return PlanGenerator(llm=OpenAIAdapter(client, model=model))
        except Exception as e:
            logging.warning(f"Failed to initialize OpenAI adapter: {e}")

    # Fallback to mock (for testing only)
    logging.warning("No LLM API key found. Using MockPlanLLM (hardcoded responses).")
    return PlanGenerator(llm=MockPlanLLM())


# ============================================================
# INSTANCES
# ============================================================


@router.get("/instances")
async def list_instances(state: Optional[str] = None):
    """List all document instances"""
    state_enum = InstanceState(state) if state else None
    instances = await instance_manager.list_instances(state=state_enum)
    return instances


@router.get("/instances/{instance_id}")
async def get_instance(instance_id: str):
    """Get instance by ID"""
    instance = await instance_manager.get_instance(instance_id)
    if not instance:
        raise HTTPException(status_code=404, detail="Instance not found")
    return instance


@router.post("/instances")
async def create_instance(request: CreateInstanceRequest = Body(...)):
    """Create a new document instance"""
    # Use configured default driver if not specified
    driver = request.driverType or config.DEFAULT_DRIVER_TYPE

    instance = await instance_manager.create_instance(
        driver_type=driver,
        document_id=request.documentId,
        file_path=request.filePath,
    )

    # Broadcast event
    await broadcast_event("instance_update", instance)

    return instance


@router.post("/instances/{instance_id}/close")
async def close_instance(instance_id: str):
    """Close an instance"""
    instance = await instance_manager.close_instance(instance_id)

    # Broadcast event
    await broadcast_event("instance_update", instance)

    return instance


@router.delete("/instances/{instance_id}")
async def delete_instance(instance_id: str):
    """Delete an instance"""
    await instance_manager.delete_instance(instance_id)
    return {"status": "deleted", "id": instance_id}


# ============================================================
# DOCUMENTS
# ============================================================


@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    title: Optional[str] = None,
    author: Optional[str] = None,
):
    """
    Upload and import a Word or PDF document

    Parses the document structure (chapters, sections, pages) and stores it in the database.
    Returns the document with its full structure.
    """
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    try:
        # Read file contents
        contents = await file.read()

        doc_format = detect_format(contents, file.filename)
        if doc_format not in (DocumentFormat.DOC, DocumentFormat.DOCX, DocumentFormat.PDF):
            raise HTTPException(
                status_code=400,
                detail="Only .doc, .docx, and .pdf files are supported",
            )

        # Create document store with database URL
        import os
        database_url = os.getenv("DATABASE_URL", "postgresql://officeplane:officeplane@db:5432/officeplane")
        doc_store = DocumentStore(database_url=database_url)

        try:
            # Lazy import to avoid requiring PIL at startup
            from ..ingestion.ingestion_service import VisionIngestionService
            from ..drivers.libreoffice_driver import LibreOfficeDriver

            ingestion_config = IngestionConfig()
            try:
                ingestion_config.validate()
            except ValueError as exc:
                raise HTTPException(status_code=400, detail=str(exc)) from exc

            # Use LibreOffice driver for DOCX->PDF conversion
            driver = LibreOfficeDriver() if doc_format in (DocumentFormat.DOC, DocumentFormat.DOCX) else None

            ingestion_service = VisionIngestionService(
                driver=driver,
                doc_store=doc_store,
                config=ingestion_config,
            )
            result = await ingestion_service.ingest(
                data=contents,
                filename=file.filename,
            )

            if not result.success or result.document is None:
                raise HTTPException(
                    status_code=500,
                    detail=f"Error ingesting document: {result.error or 'Unknown error'}",
                )

            full_document = result.document

            if not full_document:
                raise HTTPException(status_code=500, detail="Failed to load document structure")

            if title or author:
                updated = await doc_store.update_document(
                    full_document.id,
                    title=title,
                    author=author,
                )
                if updated:
                    full_document.title = updated.title
                    full_document.author = updated.author

            # Store the original file bytes for download using raw SQL
            # (Prisma Python client has issues with Bytes serialization)
            try:
                import asyncpg
                conn = await asyncpg.connect(database_url)
                try:
                    await conn.execute(
                        """
                        UPDATE documents
                        SET source_file = $1,
                            source_format = $2,
                            file_name = $3
                        WHERE id = $4
                        """,
                        contents,
                        doc_format.value,
                        file.filename,
                        full_document.id,
                    )
                finally:
                    await conn.close()
            except Exception as e:
                # Log but don't fail the upload if source file storage fails
                import logging
                logging.error(f"Failed to store source file: {e}")

            # Broadcast event
            await broadcast_event(
                "document_uploaded",
                {
                    "id": str(full_document.id),
                    "title": full_document.title,
                    "author": full_document.author,
                    "filename": file.filename,
                },
            )

            return _serialize_document(full_document)
        finally:
            await doc_store.close()

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error importing document: {str(e)}"
        )


@router.get("/documents")
async def list_documents():
    """List all documents"""
    db = await get_db()
    documents = await db.document.find_many(
        order={"createdAt": "desc"},
        include={"chapters": {"include": {"sections": True}}},
    )

    return [
        {
            "id": str(doc.id),
            "title": doc.title,
            "author": doc.author,
            "createdAt": doc.createdAt.isoformat(),
            "total_chapters": len(doc.chapters or []),
            "total_sections": sum(len(ch.sections or []) for ch in (doc.chapters or [])),
        }
        for doc in documents
    ]


@router.get("/documents/{document_id}")
async def get_document(document_id: str):
    """Get document with full structure"""
    import os
    database_url = os.getenv(
        "DATABASE_URL", "postgresql://officeplane:officeplane@db:5432/officeplane"
    )
    doc_store = DocumentStore(database_url=database_url)
    await doc_store.connect()

    try:
        full_document = await doc_store.get_document(document_id, load_children=True)

        if not full_document:
            raise HTTPException(status_code=404, detail="Document not found")

        return {
            "id": str(full_document.id),
            "title": full_document.title,
            "author": full_document.author,
            "chapters": [
                {
                    "id": str(ch.id),
                    "title": ch.title,
                    "order_index": ch.order_index,
                    "sections": [
                        {
                            "id": str(sec.id),
                            "title": sec.title,
                            "order_index": sec.order_index,
                            "page_count": len(sec.pages) if sec.pages else 0,
                        }
                        for sec in (ch.sections or [])
                    ],
                }
                for ch in (full_document.chapters or [])
            ],
            "total_chapters": len(full_document.chapters or []),
            "total_pages": sum(
                len(sec.pages or [])
                for ch in (full_document.chapters or [])
                for sec in (ch.sections or [])
            ),
        }
    finally:
        await doc_store.close()


@router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document and all its chapters, sections, and pages"""
    import os

    db = await get_db()

    # Check if document exists
    document = await db.document.find_unique(where={"id": document_id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Close any instances associated with this document
    instances = await db.documentinstance.find_many(
        where={"documentId": document_id}
    )
    for instance in instances:
        try:
            await instance_manager.close_instance(instance.id)
        except Exception:
            pass  # Instance might already be closed

    # Delete associated instances
    await db.documentinstance.delete_many(where={"documentId": document_id})

    # Delete document (cascades to chapters, sections, pages due to DB relations)
    await db.document.delete(where={"id": document_id})

    # Broadcast event
    await broadcast_event("document_deleted", {
        "id": document_id,
        "title": document.title,
    })

    return {"status": "deleted", "id": document_id}


@router.get("/documents/{document_id}/download")
async def download_document(document_id: str, format: str = "original"):
    """
    Download a document.

    Args:
        document_id: Document to download
        format: Output format - "original" (default), "docx", or "markdown"
            - "original": Returns the original uploaded file (preserves formatting)
            - "docx": Generates a new DOCX from parsed content
            - "markdown": Exports as Markdown

    Returns:
        File download response
    """
    import os
    import tempfile
    from uuid import UUID

    if format not in ("original", "docx", "markdown", "md"):
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}. Use 'original', 'docx', or 'markdown'.")

    # First, check if we have the original file stored
    db = await get_db()
    db_doc = await db.document.find_unique(where={"id": document_id})

    if not db_doc:
        raise HTTPException(status_code=404, detail="Document not found")

    safe_title = db_doc.title.replace(" ", "_").replace("/", "_")
    original_filename = db_doc.fileName or f"{safe_title}.docx"

    # Return original file if requested and available
    if format == "original":
        if db_doc.sourceFile:
            # Determine content type from source format
            content_type = "application/octet-stream"
            if db_doc.sourceFormat == "docx":
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif db_doc.sourceFormat == "pdf":
                content_type = "application/pdf"
            elif db_doc.sourceFormat == "doc":
                content_type = "application/msword"

            # Use asyncpg directly to get raw bytes (Prisma has issues with bytea)
            import os
            import asyncpg
            database_url = os.getenv(
                "DATABASE_URL", "postgresql://officeplane:officeplane@db:5432/officeplane"
            )
            conn = await asyncpg.connect(database_url)
            try:
                row = await conn.fetchrow(
                    "SELECT source_file FROM documents WHERE id = $1",
                    db_doc.id,
                )
                if row and row['source_file']:
                    file_bytes = row['source_file']
                else:
                    # Fallback to generating DOCX
                    format = "docx"
                    file_bytes = None
            finally:
                await conn.close()

            if file_bytes:
                return Response(
                    content=file_bytes,
                    media_type=content_type,
                    headers={
                        "Content-Disposition": f'attachment; filename="{original_filename}"'
                    }
                )
        else:
            # No original file, fall back to generating DOCX
            format = "docx"

    database_url = os.getenv(
        "DATABASE_URL", "postgresql://officeplane:officeplane@db:5432/officeplane"
    )
    doc_store = DocumentStore(database_url=database_url)
    await doc_store.connect()

    try:
        doc = await doc_store.get_document(UUID(document_id), load_children=True)

        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        if format in ("markdown", "md"):
            # Export as Markdown
            from officeplane.documents.exporter import DocumentExporter
            exporter = DocumentExporter(doc_store=doc_store)
            content = await exporter.export_to_markdown(UUID(document_id))

            return Response(
                content=content,
                media_type="text/markdown",
                headers={
                    "Content-Disposition": f'attachment; filename="{safe_title}.md"'
                }
            )

        else:
            # Export as DOCX (generated from parsed content)
            try:
                from docx import Document as DocxDocument
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="python-docx not installed. Install with: pip install python-docx"
                )

            # Create DOCX document
            docx = DocxDocument()

            # Title page
            title_para = docx.add_paragraph()
            title_run = title_para.add_run(doc.title)
            title_run.bold = True
            title_run.font.size = Pt(28)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if doc.author:
                author_para = docx.add_paragraph()
                author_run = author_para.add_run(f"by {doc.author}")
                author_run.font.size = Pt(14)
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            docx.add_page_break()

            # Content
            for chapter in (doc.chapters or []):
                docx.add_heading(chapter.title, level=1)

                if chapter.summary:
                    summary_para = docx.add_paragraph()
                    summary_para.add_run(chapter.summary).italic = True

                for section in (chapter.sections or []):
                    docx.add_heading(section.title, level=2)

                    for page in (section.pages or []):
                        if page.content:
                            # Split by paragraphs and add each
                            for para_text in page.content.split("\n\n"):
                                para_text = para_text.strip()
                                if para_text:
                                    docx.add_paragraph(para_text)

            # Save to bytes
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                docx.save(tmp.name)
                tmp_path = tmp.name

            try:
                with open(tmp_path, "rb") as f:
                    docx_bytes = f.read()
            finally:
                os.unlink(tmp_path)

            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="{safe_title}.docx"'
                }
            )

    finally:
        await doc_store.close()


@router.post("/documents/{document_id}/plan")
async def plan_document(document_id: str, request: PlanDocumentRequest = Body(...)):
    """Generate an action plan for editing an existing document."""
    import os
    from uuid import UUID

    database_url = os.getenv(
        "DATABASE_URL", "postgresql://officeplane:officeplane@db:5432/officeplane"
    )
    doc_store = DocumentStore(database_url=database_url)
    await doc_store.connect()

    try:
        outline = await doc_store.get_outline(UUID(document_id))
        if not outline:
            raise HTTPException(status_code=404, detail="Document not found")

        outline_text = _format_outline_for_prompt(outline)

        # Use the new edit plan generator
        generator = _build_plan_generator()
        result = await generator.generate_edit_plan(
            document_outline=outline_text,
            user_request=request.prompt,
            document_id=document_id,
            document_title=outline.title,
        )

        if not result.success:
            raise HTTPException(status_code=500, detail=result.error or "Plan generation failed")

        summary = PlanSummary.from_plan(result.plan)
        return {
            "document": {
                "id": str(outline.id),
                "title": outline.title,
                "author": outline.author,
                "chapter_count": outline.chapter_count,
                "section_count": outline.section_count,
                "page_count": outline.page_count,
            },
            "plan": summary.model_dump(),
            "tree": PlanDisplayer.to_json(result.plan, include_inputs=True),
        }
    finally:
        await doc_store.close()


class ExecutePlanRequest(BaseModel):
    """Request body for executing a plan."""
    tree: Dict[str, Any]  # The tree from plan response


class VerifyRequest(BaseModel):
    """Request body for verifying changes."""
    original_request: str  # What the user asked for
    expected_changes: Optional[List[str]] = None  # Specific things to check for


@router.post("/documents/{document_id}/execute")
async def execute_plan(document_id: str, request: ExecutePlanRequest = Body(...)):
    """
    Execute an action plan on a document.

    Takes the tree structure from the /plan response and executes each action.
    """
    import os
    from officeplane.components.planning.models import ActionNode, ActionPlan
    from officeplane.components.planning.executor import PlanExecutor

    database_url = os.getenv(
        "DATABASE_URL", "postgresql://officeplane:officeplane@db:5432/officeplane"
    )
    doc_store = DocumentStore(database_url=database_url)
    await doc_store.connect()

    try:
        # Reconstruct ActionPlan from the tree
        def build_node(node_data: Dict[str, Any]) -> ActionNode:
            children = [build_node(c) for c in node_data.get("children", [])]
            return ActionNode(
                id=node_data.get("id", f"node_{len(children)}"),
                action_name=node_data.get("action", ""),
                description=node_data.get("description", ""),
                inputs=node_data.get("inputs", {}),
                children=children,
                status=node_data.get("status", "pending"),
            )

        roots = [build_node(n) for n in request.tree.get("tree", [])]

        plan = ActionPlan(
            title=f"Execution for document {document_id}",
            original_prompt="",
            roots=roots,
        )

        # Track progress for response
        progress_log: List[Dict[str, Any]] = []

        def on_start(node: ActionNode):
            progress_log.append({
                "node_id": node.id,
                "action": node.action_name,
                "status": "running",
            })

        def on_complete(node: ActionNode, output: Dict[str, Any]):
            progress_log.append({
                "node_id": node.id,
                "action": node.action_name,
                "status": "completed",
                "output": output,
            })

        def on_failed(node: ActionNode, error: str):
            progress_log.append({
                "node_id": node.id,
                "action": node.action_name,
                "status": "failed",
                "error": error,
            })

        executor = PlanExecutor(
            doc_store=doc_store,
            on_node_start=on_start,
            on_node_complete=on_complete,
            on_node_failed=on_failed,
        )

        result = await executor.execute(plan)

        # Broadcast update
        await broadcast_event("plan_executed", {
            "document_id": document_id,
            "success": result["success"],
            "completed": result["completed"],
            "failed": result["failed"],
        })

        return {
            "success": result["success"],
            "completed": result["completed"],
            "failed": result["failed"],
            "total": result["total"],
            "progress": progress_log,
            "errors": result.get("errors", {}),
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Execution failed: {str(e)}")

    finally:
        await doc_store.close()


@router.post("/documents/{document_id}/verify")
async def verify_changes(document_id: str, request: VerifyRequest = Body(...)):
    """
    Verify that changes were applied correctly to a document.

    Uses AI to check if the document content matches the original request.
    Returns a verification report with pass/fail status and details.
    """
    import os
    from uuid import UUID

    database_url = os.getenv(
        "DATABASE_URL", "postgresql://officeplane:officeplane@db:5432/officeplane"
    )
    doc_store = DocumentStore(database_url=database_url)
    await doc_store.connect()

    try:
        # Get full document with content
        doc = await doc_store.get_document(UUID(document_id), load_children=True)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Build document content summary for verification
        content_summary = []
        for chapter in (doc.chapters or []):
            chapter_info = {
                "title": chapter.title,
                "sections": []
            }
            for section in (chapter.sections or []):
                section_info = {
                    "title": section.title,
                    "pages": []
                }
                for page in (section.pages or []):
                    section_info["pages"].append({
                        "page_number": page.page_number,
                        "content_preview": (page.content or "")[:500]
                    })
                chapter_info["sections"].append(section_info)
            content_summary.append(chapter_info)

        # Use LLM to verify changes
        generator = _build_plan_generator()

        verify_prompt = f"""You are a document verification assistant. Check if the document was updated correctly.

Original user request:
"{request.original_request}"

Current document structure and content:
{json.dumps(content_summary, indent=2)}

{f"Expected changes to verify: {request.expected_changes}" if request.expected_changes else ""}

Analyze the document and determine:
1. Was the user's request fulfilled?
2. Can you find the expected content/changes?
3. Are there any issues?

Respond with a JSON object:
{{
    "verified": true/false,
    "confidence": 0.0-1.0,
    "findings": [
        {{"check": "description of what was checked", "passed": true/false, "details": "explanation"}}
    ],
    "summary": "Brief summary of verification results",
    "suggestions": ["Any suggestions for fixes if verification failed"]
}}

Respond with ONLY the JSON, no additional text."""

        messages = [
            {"role": "user", "content": verify_prompt},
        ]

        response = await generator.llm.chat(messages, tools=[])
        content = response.get("content", "")

        if not content:
            raise HTTPException(status_code=500, detail="Verification LLM returned empty response")

        # Parse verification result
        try:
            # Handle markdown code blocks
            if "```json" in content:
                start = content.find("```json") + 7
                end = content.find("```", start)
                if end > start:
                    content = content[start:end].strip()
            elif "```" in content:
                start = content.find("```") + 3
                end = content.find("```", start)
                if end > start:
                    content = content[start:end].strip()

            verification_result = json.loads(content)
        except json.JSONDecodeError as e:
            verification_result = {
                "verified": False,
                "confidence": 0.0,
                "findings": [{"check": "JSON parsing", "passed": False, "details": f"Failed to parse LLM response: {e}"}],
                "summary": "Verification could not be completed due to parsing error",
                "suggestions": ["Try running verification again"],
                "raw_response": content[:500]
            }

        return {
            "document_id": document_id,
            "document_title": doc.title,
            "original_request": request.original_request,
            "verification": verification_result,
            "document_stats": {
                "chapters": len(doc.chapters or []),
                "sections": sum(len(ch.sections or []) for ch in (doc.chapters or [])),
                "pages": sum(
                    len(sec.pages or [])
                    for ch in (doc.chapters or [])
                    for sec in (ch.sections or [])
                ),
            }
        }

    finally:
        await doc_store.close()


# ============================================================
# TASKS
# ============================================================


@router.get("/tasks")
async def list_tasks(
    state: Optional[str] = None,
    priority: Optional[str] = None,
    limit: int = 100,
):
    """List tasks"""
    state_enum = TaskState(state) if state else None
    priority_enum = TaskPriority(priority) if priority else None

    tasks = await task_queue.list_tasks(
        state=state_enum,
        priority=priority_enum,
        limit=limit,
    )
    return tasks


@router.get("/tasks/{task_id}")
async def get_task(task_id: str):
    """Get task by ID"""
    task = await task_queue.get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return task


@router.post("/tasks")
async def create_task(
    taskType: str,
    payload: dict = {},
    taskName: Optional[str] = None,
    documentId: Optional[str] = None,
    instanceId: Optional[str] = None,
    priority: str = "NORMAL",
    maxRetries: int = 3,
):
    """Enqueue a new task"""
    priority_enum = TaskPriority(priority)

    task = await task_queue.enqueue_task(
        task_type=taskType,
        payload=payload,
        task_name=taskName,
        document_id=documentId,
        instance_id=instanceId,
        priority=priority_enum,
        max_retries=maxRetries,
    )

    # Broadcast event
    await broadcast_event("task_update", task)

    return task


@router.post("/tasks/{task_id}/cancel")
async def cancel_task(task_id: str):
    """Cancel a task"""
    task = await task_queue.cancel_task(task_id)

    # Broadcast event
    await broadcast_event("task_update", task)

    return task


@router.post("/tasks/{task_id}/retry")
async def retry_task(task_id: str):
    """Retry a failed task"""
    task = await task_queue.retry_task(task_id)

    # Broadcast event
    await broadcast_event("task_update", task)

    return task


# ============================================================
# HISTORY
# ============================================================


@router.get("/history")
async def get_history(
    eventType: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
):
    """Get execution history"""
    db = await get_db()

    where_clause = {}
    if eventType:
        where_clause["eventType"] = EventType(eventType)

    history = await db.executionhistory.find_many(
        where=where_clause,
        include={"document": True, "task": True},
        order={"timestamp": "desc"},
        skip=offset,
        take=limit,
    )

    return [event.model_dump() for event in history]


# ============================================================
# METRICS
# ============================================================


@router.get("/metrics")
async def get_metrics():
    """Get system metrics"""
    db = await get_db()

    # Instance metrics
    instances = await db.documentinstance.find_many()
    instances_by_state = {}
    total_memory = 0
    total_cpu = 0
    active_count = 0

    for inst in instances:
        state = inst.state  # state is already a string from the database
        instances_by_state[state] = instances_by_state.get(state, 0) + 1

        if state in [InstanceState.OPEN.value, InstanceState.IDLE.value, InstanceState.IN_USE.value]:
            active_count += 1
            if inst.memoryMb:
                total_memory += inst.memoryMb
            if inst.cpuPercent:
                total_cpu += inst.cpuPercent

    # Task metrics
    tasks = await db.task.find_many()
    tasks_by_state = {}
    total_duration = 0
    completed_count = 0
    failed_count = 0

    for task in tasks:
        state = task.state  # state is already a string from the database
        tasks_by_state[state] = tasks_by_state.get(state, 0) + 1

        if state == TaskState.COMPLETED.value and task.startedAt and task.completedAt:
            duration = (task.completedAt - task.startedAt).total_seconds() * 1000
            total_duration += duration
            completed_count += 1

        if state == TaskState.FAILED.value:
            failed_count += 1

    avg_duration_ms = total_duration / completed_count if completed_count > 0 else 0
    failure_rate = failed_count / len(tasks) if tasks else 0

    return {
        "instances": {
            "total": len(instances),
            "byState": instances_by_state,
        },
        "tasks": {
            "total": len(tasks),
            "byState": tasks_by_state,
            "avgDurationMs": avg_duration_ms,
            "failureRate": failure_rate,
        },
        "system": {
            "uptime": 0,  # TODO: track uptime
            "memoryUsageMb": total_memory / active_count if active_count > 0 else 0,
            "cpuPercent": total_cpu / active_count if active_count > 0 else 0,
        },
    }


# ============================================================
# WEBSOCKET
# ============================================================


@router.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    """WebSocket endpoint for real-time updates"""
    await websocket.accept()
    websocket_connections.add(websocket)

    print(f"[WebSocket] Client connected. Total connections: {len(websocket_connections)}")

    try:
        # Send initial connection message
        await websocket.send_json({
            "type": "connected",
            "data": {"message": "Connected to OfficePlane Management System"},
            "timestamp": datetime.utcnow().isoformat(),
        })

        # Keep connection alive and listen for messages
        while True:
            data = await websocket.receive_text()
            # Echo or process messages if needed
            # await websocket.send_json({"type": "echo", "data": data})

    except WebSocketDisconnect:
        websocket_connections.discard(websocket)
        print(
            f"[WebSocket] Client disconnected. Total connections: {len(websocket_connections)}"
        )
    except Exception as e:
        print(f"[WebSocket] Error: {e}")
        websocket_connections.discard(websocket)
