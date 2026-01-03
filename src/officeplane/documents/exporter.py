"""
Document exporter for DOCX and PDF formats.

Converts the document hierarchy to output files using:
- python-docx for DOCX generation
- Existing OfficePlane render pipeline for PDF conversion
"""

from __future__ import annotations

import logging
import os
import tempfile
from typing import Optional
from uuid import UUID

from officeplane.components.context import ComponentContext
from officeplane.documents.store import DocumentStore

log = logging.getLogger("officeplane.documents.exporter")


class DocumentExporter:
    """
    Export documents to DOCX and PDF formats.

    Usage:
        exporter = DocumentExporter(doc_store=store)
        url = await exporter.export_to_docx(document_id, ctx)
        url = await exporter.export_to_pdf(document_id, ctx)
    """

    def __init__(self, doc_store: Optional[DocumentStore] = None) -> None:
        self._doc_store = doc_store

    async def _get_store(self) -> DocumentStore:
        """Get or create the document store."""
        if self._doc_store is None:
            self._doc_store = DocumentStore()
            await self._doc_store.connect()
        return self._doc_store

    async def export_to_docx(
        self,
        document_id: UUID,
        ctx: ComponentContext,
        output_path: Optional[str] = None,
    ) -> str:
        """
        Export a document to DOCX format.

        Args:
            document_id: Document to export
            ctx: Component context (for artifact storage)
            output_path: Optional file path (if not provided, stores as artifact)

        Returns:
            File URL or path
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        store = await self._get_store()
        doc = await store.get_document(document_id, load_children=True)

        if not doc:
            raise ValueError(f"Document not found: {document_id}")

        # Create DOCX document
        docx = DocxDocument()

        # Title page
        title_para = docx.add_paragraph()
        title_run = title_para.add_run(doc.title)
        title_run.bold = True
        title_run.font.size = Pt(28)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if doc.author:
            author_para = docx.add_paragraph()
            author_run = author_para.add_run(f"by {doc.author}")
            author_run.font.size = Pt(14)
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        docx.add_page_break()

        # Table of Contents placeholder
        docx.add_heading("Table of Contents", level=1)
        docx.add_paragraph("(Table of contents will be auto-generated in Word)")
        docx.add_page_break()

        # Content
        for chapter in doc.chapters:
            # Chapter heading
            docx.add_heading(chapter.title, level=1)

            if chapter.summary:
                summary_para = docx.add_paragraph()
                summary_para.add_run(chapter.summary).italic = True

            for section in chapter.sections:
                # Section heading
                docx.add_heading(section.title, level=2)

                for page in section.pages:
                    if page.content:
                        # Parse markdown-ish content to paragraphs
                        self._add_content_to_docx(docx, page.content)

        # Save
        if output_path:
            docx.save(output_path)
            log.info(f"Exported DOCX to: {output_path}")
            return output_path
        else:
            # Save to temp file and store as artifact
            with tempfile.NamedTemporaryFile(
                suffix=".docx", delete=False
            ) as tmp:
                docx.save(tmp.name)
                tmp_path = tmp.name

            try:
                # Read and store as artifact
                with open(tmp_path, "rb") as f:
                    data = f.read()

                if ctx.store:
                    filename = f"{doc.title.replace(' ', '_')}.docx"
                    url = ctx.store.put_bytes(
                        request_id=ctx.request_id,
                        name=filename,
                        data=data,
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    log.info(f"Stored DOCX as artifact: {url}")
                    return url
                else:
                    # No store available, keep temp file
                    log.info(f"Exported DOCX to temp: {tmp_path}")
                    return tmp_path
            finally:
                # Clean up if we stored it
                if ctx.store and os.path.exists(tmp_path):
                    os.unlink(tmp_path)

    def _add_content_to_docx(self, docx: object, content: str) -> None:
        """Add markdown-ish content to DOCX document."""
        # Simple paragraph splitting
        paragraphs = content.split("\n\n")

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check for headings (### Header)
            if para_text.startswith("### "):
                docx.add_heading(para_text[4:], level=3)  # type: ignore
            elif para_text.startswith("## "):
                docx.add_heading(para_text[3:], level=2)  # type: ignore
            elif para_text.startswith("# "):
                docx.add_heading(para_text[2:], level=1)  # type: ignore
            elif para_text.startswith("- ") or para_text.startswith("* "):
                # Bullet list items
                lines = para_text.split("\n")
                for line in lines:
                    line = line.strip()
                    if line.startswith("- ") or line.startswith("* "):
                        docx.add_paragraph(line[2:], style="List Bullet")  # type: ignore
                    elif line:
                        docx.add_paragraph(line)  # type: ignore
            else:
                # Regular paragraph
                para = docx.add_paragraph()  # type: ignore
                # Handle inline bold/italic
                self._add_formatted_text(para, para_text)

    def _add_formatted_text(self, para: object, text: str) -> None:
        """Add text with basic markdown formatting (bold, italic)."""
        import re

        # Simple pattern matching for **bold** and *italic*
        # This is a basic implementation
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)

        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])  # type: ignore
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])  # type: ignore
                run.italic = True
            else:
                para.add_run(part)  # type: ignore

    async def export_to_pdf(
        self,
        document_id: UUID,
        ctx: ComponentContext,
        output_path: Optional[str] = None,
    ) -> str:
        """
        Export a document to PDF format.

        Uses the DOCX export + OfficePlane driver for conversion.

        Args:
            document_id: Document to export
            ctx: Component context (for driver and artifact storage)
            output_path: Optional file path

        Returns:
            File URL or path
        """
        # First export to DOCX
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_path = tmp.name

        try:
            await self.export_to_docx(document_id, ctx, output_path=docx_path)

            # Convert to PDF using driver
            if ctx.driver:
                with open(docx_path, "rb") as f:
                    docx_bytes = f.read()

                store = await self._get_store()
                doc = await store.get_document(document_id)
                base_filename = doc.title.replace(" ", "_") if doc else "document"

                pdf_bytes = ctx.driver.convert_to_pdf(
                    f"{base_filename}.docx", docx_bytes
                )

                if output_path:
                    with open(output_path, "wb") as f:
                        f.write(pdf_bytes)
                    log.info(f"Exported PDF to: {output_path}")
                    return output_path
                elif ctx.store:
                    filename = f"{base_filename}.pdf"
                    url = ctx.store.put_bytes(
                        request_id=ctx.request_id,
                        name=filename,
                        data=pdf_bytes,
                        content_type="application/pdf",
                    )
                    log.info(f"Stored PDF as artifact: {url}")
                    return url
                else:
                    # Save to temp file
                    with tempfile.NamedTemporaryFile(
                        suffix=".pdf", delete=False
                    ) as pdf_tmp:
                        pdf_tmp.write(pdf_bytes)
                        log.info(f"Exported PDF to temp: {pdf_tmp.name}")
                        return pdf_tmp.name
            else:
                raise RuntimeError(
                    "No driver available for PDF conversion. "
                    "Ensure ComponentContext has a driver configured."
                )
        finally:
            if os.path.exists(docx_path):
                os.unlink(docx_path)

    async def export_to_markdown(
        self,
        document_id: UUID,
        output_path: Optional[str] = None,
    ) -> str:
        """
        Export a document to Markdown format.

        Args:
            document_id: Document to export
            output_path: Optional file path

        Returns:
            Markdown content or file path
        """
        store = await self._get_store()
        doc = await store.get_document(document_id, load_children=True)

        if not doc:
            raise ValueError(f"Document not found: {document_id}")

        lines = []

        # Title
        lines.append(f"# {doc.title}")
        lines.append("")
        if doc.author:
            lines.append(f"*by {doc.author}*")
            lines.append("")
        lines.append("---")
        lines.append("")

        # Table of Contents
        lines.append("## Table of Contents")
        lines.append("")
        for chapter in doc.chapters:
            lines.append(f"- [{chapter.title}](#{self._slugify(chapter.title)})")
            for section in chapter.sections:
                lines.append(
                    f"  - [{section.title}](#{self._slugify(section.title)})"
                )
        lines.append("")
        lines.append("---")
        lines.append("")

        # Content
        for chapter in doc.chapters:
            lines.append(f"## {chapter.title}")
            lines.append("")
            if chapter.summary:
                lines.append(f"*{chapter.summary}*")
                lines.append("")

            for section in chapter.sections:
                lines.append(f"### {section.title}")
                lines.append("")

                for page in section.pages:
                    if page.content:
                        lines.append(page.content)
                        lines.append("")

        markdown = "\n".join(lines)

        if output_path:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(markdown)
            log.info(f"Exported Markdown to: {output_path}")
            return output_path
        else:
            return markdown

    def _slugify(self, text: str) -> str:
        """Convert text to URL-safe slug."""
        import re

        slug = text.lower()
        slug = re.sub(r"[^\w\s-]", "", slug)
        slug = re.sub(r"[\s_]+", "-", slug)
        return slug.strip("-")
