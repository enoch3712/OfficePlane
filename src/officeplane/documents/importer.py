"""
Document importer for Word (.docx) files.

Parses Word documents and maps them to the document hierarchy:
- Heading 1 → Chapter
- Heading 2 → Section
- Body text → Page content

Supports round-trip editing: import → modify → export.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional
from uuid import UUID

from officeplane.documents.store import DocumentStore
from officeplane.documents.models import DocumentModel

log = logging.getLogger("officeplane.documents.importer")


@dataclass
class ParsedParagraph:
    """A parsed paragraph from Word."""

    text: str
    style: str  # e.g., "Heading 1", "Normal", "List Bullet"
    is_heading: bool = False
    heading_level: int = 0  # 1 for Heading 1, 2 for Heading 2, etc.


@dataclass
class ParsedSection:
    """A parsed section (Heading 2 level)."""

    title: str
    paragraphs: List[str] = field(default_factory=list)


@dataclass
class ParsedChapter:
    """A parsed chapter (Heading 1 level)."""

    title: str
    sections: List[ParsedSection] = field(default_factory=list)
    # Content before first section
    preamble: List[str] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """A fully parsed document structure."""

    title: str
    author: Optional[str] = None
    chapters: List[ParsedChapter] = field(default_factory=list)
    # Content before first chapter
    preamble: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentImporter:
    """
    Import Word documents into the OfficePlane document model.

    Parses document structure based on heading styles:
    - Heading 1 → Chapter
    - Heading 2 → Section
    - Normal/Body text → Page content

    Usage:
        importer = DocumentImporter(doc_store=store)
        await importer.connect()

        # Import from file path
        doc = await importer.import_from_file("/path/to/document.docx")

        # Import from bytes
        doc = await importer.import_from_bytes(docx_bytes, title="My Doc")

        # Now use AuthorComponent to modify, then export back
    """

    def __init__(
        self,
        doc_store: Optional[DocumentStore] = None,
        words_per_page: int = 500,
    ) -> None:
        """
        Initialize the importer.

        Args:
            doc_store: DocumentStore for persistence
            words_per_page: Target words per page when splitting content
        """
        self._doc_store = doc_store
        self.words_per_page = words_per_page

    async def _get_store(self) -> DocumentStore:
        """Get or create document store."""
        if self._doc_store is None:
            self._doc_store = DocumentStore()
            await self._doc_store.connect()
        return self._doc_store

    async def connect(self) -> None:
        """Connect to the document store."""
        await self._get_store()

    async def close(self) -> None:
        """Close connections."""
        if self._doc_store:
            await self._doc_store.close()

    def parse_docx(self, docx_bytes: bytes) -> ParsedDocument:
        """
        Parse a Word document into structured format.

        Args:
            docx_bytes: Raw bytes of the .docx file

        Returns:
            ParsedDocument with chapters, sections, and content
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        doc = DocxDocument(io.BytesIO(docx_bytes))

        # Extract document properties
        title = "Untitled Document"
        author = None

        if doc.core_properties:
            if doc.core_properties.title:
                title = doc.core_properties.title
            if doc.core_properties.author:
                author = doc.core_properties.author

        # Parse all paragraphs
        paragraphs = self._parse_paragraphs(doc)

        # Build document structure from paragraphs
        parsed = self._build_structure(paragraphs, title, author)

        return parsed

    def _parse_paragraphs(self, doc: Any) -> List[ParsedParagraph]:
        """Extract paragraphs with their styles."""
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"

            # Detect heading level
            is_heading = False
            heading_level = 0

            if style_name.startswith("Heading"):
                is_heading = True
                # Extract level number (e.g., "Heading 1" → 1)
                match = re.search(r"Heading\s*(\d+)", style_name)
                if match:
                    heading_level = int(match.group(1))

            paragraphs.append(
                ParsedParagraph(
                    text=text,
                    style=style_name,
                    is_heading=is_heading,
                    heading_level=heading_level,
                )
            )

        return paragraphs

    def _build_structure(
        self,
        paragraphs: List[ParsedParagraph],
        title: str,
        author: Optional[str],
    ) -> ParsedDocument:
        """Build document structure from parsed paragraphs."""
        parsed = ParsedDocument(title=title, author=author)

        current_chapter: Optional[ParsedChapter] = None
        current_section: Optional[ParsedSection] = None

        for para in paragraphs:
            if para.is_heading and para.heading_level == 1:
                # New chapter
                current_chapter = ParsedChapter(title=para.text)
                parsed.chapters.append(current_chapter)
                current_section = None

            elif para.is_heading and para.heading_level == 2:
                # New section
                current_section = ParsedSection(title=para.text)
                if current_chapter:
                    current_chapter.sections.append(current_section)
                else:
                    # Section without chapter - create default chapter
                    current_chapter = ParsedChapter(title="Main Content")
                    parsed.chapters.append(current_chapter)
                    current_chapter.sections.append(current_section)

            else:
                # Body content
                if current_section:
                    current_section.paragraphs.append(para.text)
                elif current_chapter:
                    current_chapter.preamble.append(para.text)
                else:
                    parsed.preamble.append(para.text)

        # Handle documents with no structure
        if not parsed.chapters:
            # Create single chapter/section for flat documents
            all_content = parsed.preamble
            if all_content:
                chapter = ParsedChapter(title="Content")
                section = ParsedSection(title="Main")
                section.paragraphs = all_content
                chapter.sections.append(section)
                parsed.chapters.append(chapter)
                parsed.preamble = []

        return parsed

    def _split_into_pages(self, paragraphs: List[str]) -> List[str]:
        """
        Split paragraphs into page-sized chunks.

        Args:
            paragraphs: List of paragraph texts

        Returns:
            List of page contents (joined paragraphs)
        """
        if not paragraphs:
            return []

        pages = []
        current_page: List[str] = []
        current_words = 0

        for para in paragraphs:
            para_words = len(para.split())

            if current_words + para_words > self.words_per_page and current_page:
                # Start new page
                pages.append("\n\n".join(current_page))
                current_page = [para]
                current_words = para_words
            else:
                current_page.append(para)
                current_words += para_words

        # Add remaining content
        if current_page:
            pages.append("\n\n".join(current_page))

        return pages

    async def import_from_bytes(
        self,
        docx_bytes: bytes,
        title: Optional[str] = None,
        author: Optional[str] = None,
        index_for_search: bool = True,
    ) -> DocumentModel:
        """
        Import a Word document from bytes.

        Args:
            docx_bytes: Raw bytes of the .docx file
            title: Override document title (uses file title if not provided)
            author: Override author
            index_for_search: Whether to index content for RAG search

        Returns:
            The created DocumentModel
        """
        store = await self._get_store()

        # Parse the document
        parsed = self.parse_docx(docx_bytes)

        # Override title/author if provided
        doc_title = title or parsed.title
        doc_author = author or parsed.author

        # Create the document
        document = await store.create_document(
            title=doc_title,
            author=doc_author,
            metadata={"imported": True, "source_format": "docx"},
        )

        log.info(f"Created document: {document.id} - {doc_title}")

        # Import chapters
        for chapter_idx, parsed_chapter in enumerate(parsed.chapters):
            chapter = await store.create_chapter(
                document_id=document.id,
                title=parsed_chapter.title,
                order_index=chapter_idx,
            )
            log.info(f"  Created chapter: {chapter.title}")

            # Handle chapter preamble as a section
            if parsed_chapter.preamble:
                preamble_section = await store.create_section(
                    chapter_id=chapter.id,
                    title="Introduction",
                    order_index=0,
                )
                pages = self._split_into_pages(parsed_chapter.preamble)
                for page_num, content in enumerate(pages, start=1):
                    await store.create_page(
                        section_id=preamble_section.id,
                        content=content,
                        page_number=page_num,
                    )

            # Import sections
            section_offset = 1 if parsed_chapter.preamble else 0
            for section_idx, parsed_section in enumerate(parsed_chapter.sections):
                section = await store.create_section(
                    chapter_id=chapter.id,
                    title=parsed_section.title,
                    order_index=section_idx + section_offset,
                )
                log.info(f"    Created section: {section.title}")

                # Split section content into pages
                pages = self._split_into_pages(parsed_section.paragraphs)
                for page_num, content in enumerate(pages, start=1):
                    await store.create_page(
                        section_id=section.id,
                        content=content,
                        page_number=page_num,
                    )
                    log.debug(f"      Created page {page_num}: {len(content)} chars")

        # Index for search if requested
        if index_for_search:
            await self._index_document(document.id)

        # Return the full document with children loaded
        result = await store.get_document(document.id, load_children=True)
        return result  # type: ignore

    async def import_from_file(
        self,
        file_path: str,
        title: Optional[str] = None,
        author: Optional[str] = None,
        index_for_search: bool = True,
    ) -> DocumentModel:
        """
        Import a Word document from a file path.

        Args:
            file_path: Path to the .docx file
            title: Override document title
            author: Override author
            index_for_search: Whether to index for RAG search

        Returns:
            The created DocumentModel
        """
        import os

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        with open(file_path, "rb") as f:
            docx_bytes = f.read()

        # Use filename as default title
        if title is None:
            title = os.path.splitext(os.path.basename(file_path))[0]

        return await self.import_from_bytes(
            docx_bytes=docx_bytes,
            title=title,
            author=author,
            index_for_search=index_for_search,
        )

    async def _index_document(self, document_id: UUID) -> int:
        """
        Index all pages in a document for RAG search.

        Args:
            document_id: Document to index

        Returns:
            Number of chunks created
        """
        try:
            from officeplane.memory.rag import DocumentIndexer
            from officeplane.memory.vector_store import VectorStore
            from officeplane.memory.embeddings import EmbeddingClient

            indexer = DocumentIndexer(
                vector_store=VectorStore(),
                embedding_client=EmbeddingClient(),
            )
            await indexer.connect()

            store = await self._get_store()
            doc = await store.get_document(document_id, load_children=True)

            if not doc:
                return 0

            total_chunks = 0
            for chapter in doc.chapters:
                for section in chapter.sections:
                    for page in section.pages:
                        if page.content:
                            chunk_ids = await indexer.index_page(
                                page_id=page.id,
                                document_id=doc.id,
                                chapter_id=chapter.id,
                                section_id=section.id,
                                content=page.content,
                            )
                            total_chunks += len(chunk_ids)

            await indexer.close()
            log.info(f"Indexed {total_chunks} chunks for document {document_id}")
            return total_chunks

        except Exception as e:
            log.warning(f"Failed to index document: {e}")
            return 0

    async def reimport_document(
        self,
        document_id: UUID,
        docx_bytes: bytes,
        preserve_structure: bool = True,
    ) -> DocumentModel:
        """
        Re-import a document, updating content from new Word file.

        This is useful for round-trip editing:
        1. Export document to Word
        2. Edit in Word
        3. Re-import to update content

        Args:
            document_id: Existing document to update
            docx_bytes: New Word document bytes
            preserve_structure: If True, tries to match chapters/sections by title

        Returns:
            Updated DocumentModel
        """
        store = await self._get_store()

        # Get existing document
        existing = await store.get_document(document_id, load_children=True)
        if not existing:
            raise ValueError(f"Document not found: {document_id}")

        # Parse new content
        parsed = self.parse_docx(docx_bytes)

        if preserve_structure:
            await self._update_matching_structure(existing, parsed)
        else:
            # Delete all existing content and re-import
            await store.delete_document(document_id)
            return await self.import_from_bytes(
                docx_bytes,
                title=existing.title,
                author=existing.author,
            )

        # Reload and return
        result = await store.get_document(document_id, load_children=True)
        return result  # type: ignore

    async def _update_matching_structure(
        self,
        existing: DocumentModel,
        parsed: ParsedDocument,
    ) -> None:
        """Update document by matching chapters/sections by title."""
        store = await self._get_store()

        # Build lookup by title
        existing_chapters = {ch.title: ch for ch in existing.chapters}

        for parsed_chapter in parsed.chapters:
            if parsed_chapter.title in existing_chapters:
                # Update existing chapter
                existing_ch = existing_chapters[parsed_chapter.title]
                existing_sections = {s.title: s for s in existing_ch.sections}

                for parsed_section in parsed_chapter.sections:
                    if parsed_section.title in existing_sections:
                        # Update existing section's pages
                        existing_sec = existing_sections[parsed_section.title]
                        pages = self._split_into_pages(parsed_section.paragraphs)

                        # Delete old pages
                        for page in existing_sec.pages:
                            await store.delete_page(page.id)

                        # Create new pages
                        for page_num, content in enumerate(pages, start=1):
                            await store.create_page(
                                section_id=existing_sec.id,
                                content=content,
                                page_number=page_num,
                            )
                    else:
                        # New section
                        section = await store.create_section(
                            chapter_id=existing_ch.id,
                            title=parsed_section.title,
                        )
                        pages = self._split_into_pages(parsed_section.paragraphs)
                        for page_num, content in enumerate(pages, start=1):
                            await store.create_page(
                                section_id=section.id,
                                content=content,
                                page_number=page_num,
                            )
            else:
                # New chapter
                chapter = await store.create_chapter(
                    document_id=existing.id,
                    title=parsed_chapter.title,
                )
                for parsed_section in parsed_chapter.sections:
                    section = await store.create_section(
                        chapter_id=chapter.id,
                        title=parsed_section.title,
                    )
                    pages = self._split_into_pages(parsed_section.paragraphs)
                    for page_num, content in enumerate(pages, start=1):
                        await store.create_page(
                            section_id=section.id,
                            content=content,
                            page_number=page_num,
                        )


class DocumentEditor:
    """
    Edit Word documents in-place, preserving all original formatting and styles.

    This class allows inserting new content into an existing document while
    keeping the original styling intact. It copies formatting from existing
    paragraphs to ensure new content matches the document's style.

    Usage:
        editor = DocumentEditor(original_docx_bytes)

        # Insert a section after paragraph 10
        editor.insert_section_after_paragraph(
            paragraph_index=10,
            heading="New Section",
            paragraphs=["First paragraph.", "Second paragraph."],
        )

        # Get the modified document
        modified_bytes = editor.save()
    """

    def __init__(self, docx_bytes: bytes) -> None:
        """
        Initialize editor with an existing document.

        Args:
            docx_bytes: Raw bytes of the original .docx file
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        self._docx_bytes = docx_bytes
        self._doc = DocxDocument(io.BytesIO(docx_bytes))
        self._body_style_para: Optional[Any] = None
        self._heading_style_paras: Dict[int, Any] = {}
        self._analyze_styles()

    def _analyze_styles(self) -> None:
        """Analyze the document to find reference paragraphs for each style."""
        for para in self._doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Check if it's a heading
            if style_name.startswith("Heading"):
                match = re.search(r"Heading\s*(\d+)", style_name)
                if match:
                    level = int(match.group(1))
                    if level not in self._heading_style_paras:
                        self._heading_style_paras[level] = para

            # Find a good body text paragraph (not a heading, list, or title)
            elif self._body_style_para is None:
                # Skip very short paragraphs, likely titles or labels
                if len(text) > 20 and not style_name.startswith(("Title", "Heading", "TOC")):
                    self._body_style_para = para

    def _get_nearest_heading_ref(self, requested_level: int) -> Optional[Any]:
        """
        Get the nearest available heading reference paragraph.

        If the exact heading level doesn't exist, find the closest one.
        """
        if not self._heading_style_paras:
            return None

        # Try exact match first
        if requested_level in self._heading_style_paras:
            return self._heading_style_paras[requested_level]

        # Find nearest available level
        available_levels = sorted(self._heading_style_paras.keys())

        # Find closest level
        closest_level = min(available_levels, key=lambda x: abs(x - requested_level))
        return self._heading_style_paras[closest_level]

    def get_available_heading_levels(self) -> List[int]:
        """Get the heading levels available in this document."""
        return sorted(self._heading_style_paras.keys())

    def _copy_paragraph_format(self, source_para: Any, target_para: Any) -> None:
        """Copy paragraph formatting from source to target."""
        from copy import deepcopy
        from docx.oxml.ns import qn

        # Copy paragraph properties (pPr) if they exist
        source_pPr = source_para._element.pPr
        if source_pPr is not None:
            # Remove existing pPr from target
            target_pPr = target_para._element.pPr
            if target_pPr is not None:
                target_para._element.remove(target_pPr)
            # Deep copy and insert the source pPr
            new_pPr = deepcopy(source_pPr)
            target_para._element.insert(0, new_pPr)

    def _copy_run_format(self, source_run: Any, target_run: Any) -> None:
        """Copy run formatting from source to target."""
        from copy import deepcopy

        # Copy run properties (rPr) if they exist
        source_rPr = source_run._element.rPr
        if source_rPr is not None:
            # Remove existing rPr from target
            target_rPr = target_run._element.rPr
            if target_rPr is not None:
                target_run._element.remove(target_rPr)
            # Deep copy and insert the source rPr
            new_rPr = deepcopy(source_rPr)
            target_run._element.insert(0, new_rPr)

    def _create_styled_paragraph(
        self,
        text: str,
        reference_para: Optional[Any] = None,
    ) -> Any:
        """Create a new paragraph with formatting copied from a reference."""
        new_para = self._doc.add_paragraph()

        if reference_para is not None:
            # Copy paragraph-level formatting
            self._copy_paragraph_format(reference_para, new_para)

            # Create a run with the text
            new_run = new_para.add_run(text)

            # Copy run formatting from the first run of the reference
            if reference_para.runs:
                self._copy_run_format(reference_para.runs[0], new_run)
        else:
            new_para.add_run(text)

        return new_para

    def get_paragraph_style_info(self, paragraph_index: int) -> Dict[str, Any]:
        """
        Get detailed style information for a paragraph.

        Args:
            paragraph_index: Index of the paragraph

        Returns:
            Dictionary with style information
        """
        if paragraph_index < 0 or paragraph_index >= len(self._doc.paragraphs):
            raise IndexError(f"Paragraph index {paragraph_index} out of range")

        para = self._doc.paragraphs[paragraph_index]
        info: Dict[str, Any] = {
            "style_name": para.style.name if para.style else None,
            "text_preview": para.text[:50] if para.text else "",
            "runs": [],
        }

        for run in para.runs:
            run_info = {
                "text": run.text[:20] if run.text else "",
                "bold": run.bold,
                "italic": run.italic,
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
            }
            info["runs"].append(run_info)

        return info

    @property
    def paragraph_count(self) -> int:
        """Get the number of paragraphs in the document."""
        return len(self._doc.paragraphs)

    def get_paragraph_texts(self) -> List[str]:
        """Get all paragraph texts for inspection."""
        return [p.text for p in self._doc.paragraphs]

    def find_paragraph_index(self, text: str, exact: bool = False) -> Optional[int]:
        """
        Find the index of a paragraph containing the given text.

        Args:
            text: Text to search for
            exact: If True, match exact text; if False, match substring

        Returns:
            Paragraph index or None if not found
        """
        for i, para in enumerate(self._doc.paragraphs):
            if exact and para.text == text:
                return i
            elif not exact and text in para.text:
                return i
        return None

    def insert_paragraph_after(
        self,
        paragraph_index: int,
        text: str,
        copy_style_from: Optional[int] = None,
    ) -> None:
        """
        Insert a new paragraph after the specified paragraph.

        Args:
            paragraph_index: Index of paragraph to insert after
            text: Text for the new paragraph
            copy_style_from: Index of paragraph to copy style from (default: nearby body text)
        """
        if paragraph_index < 0 or paragraph_index >= len(self._doc.paragraphs):
            raise IndexError(f"Paragraph index {paragraph_index} out of range")

        # Get the reference paragraph for positioning
        ref_para = self._doc.paragraphs[paragraph_index]

        # Determine style source
        if copy_style_from is not None:
            style_source = self._doc.paragraphs[copy_style_from]
        elif self._body_style_para is not None:
            style_source = self._body_style_para
        else:
            style_source = ref_para

        # Create styled paragraph
        new_para = self._create_styled_paragraph(text, style_source)

        # Move the new paragraph after the reference paragraph
        ref_para._element.addnext(new_para._element)

    def insert_section_after_paragraph(
        self,
        paragraph_index: int,
        heading: str,
        paragraphs: List[str],
        heading_level: int = 2,
    ) -> int:
        """
        Insert a complete section (heading + paragraphs) after a paragraph.

        The heading and body text will copy formatting from existing paragraphs
        in the document to maintain consistent styling.

        Args:
            paragraph_index: Index of paragraph to insert after
            heading: Section heading text
            paragraphs: List of paragraph texts
            heading_level: Heading level (1-9)

        Returns:
            Index of the last inserted paragraph
        """
        if paragraph_index < 0 or paragraph_index >= len(self._doc.paragraphs):
            raise IndexError(f"Paragraph index {paragraph_index} out of range")

        # Get the reference paragraph's XML element
        ref_para = self._doc.paragraphs[paragraph_index]
        current_element = ref_para._element

        # Find a heading to copy style from (uses nearest available if exact level not found)
        heading_ref = self._get_nearest_heading_ref(heading_level)

        # Insert heading with copied style
        if heading_ref is not None:
            heading_para = self._create_styled_paragraph(heading, heading_ref)
        else:
            # Fallback: create with built-in style (shouldn't happen if document has any headings)
            heading_para = self._doc.add_paragraph(heading)
            try:
                heading_para.style = f"Heading {heading_level}"
            except KeyError:
                pass  # Style doesn't exist, use default

        current_element.addnext(heading_para._element)
        current_element = heading_para._element

        # Insert each paragraph with copied body style
        body_ref = self._body_style_para
        for para_text in paragraphs:
            if body_ref is not None:
                new_para = self._create_styled_paragraph(para_text, body_ref)
            else:
                new_para = self._doc.add_paragraph(para_text)

            current_element.addnext(new_para._element)
            current_element = new_para._element

        return paragraph_index + 1 + len(paragraphs)

    def insert_paragraphs_after(
        self,
        paragraph_index: int,
        paragraphs: List[str],
        copy_style_from: Optional[int] = None,
    ) -> int:
        """
        Insert multiple paragraphs after a specified paragraph.

        Args:
            paragraph_index: Index of paragraph to insert after
            paragraphs: List of paragraph texts
            copy_style_from: Index of paragraph to copy style from (default: auto-detect body style)

        Returns:
            Index of the last inserted paragraph
        """
        if paragraph_index < 0 or paragraph_index >= len(self._doc.paragraphs):
            raise IndexError(f"Paragraph index {paragraph_index} out of range")

        ref_para = self._doc.paragraphs[paragraph_index]
        current_element = ref_para._element

        # Determine style source
        if copy_style_from is not None:
            style_source = self._doc.paragraphs[copy_style_from]
        else:
            style_source = self._body_style_para

        for para_text in paragraphs:
            if style_source is not None:
                new_para = self._create_styled_paragraph(para_text, style_source)
            else:
                new_para = self._doc.add_paragraph(para_text)

            current_element.addnext(new_para._element)
            current_element = new_para._element

        return paragraph_index + len(paragraphs)

    def insert_at_middle(
        self,
        heading: str,
        paragraphs: List[str],
        heading_level: int = 2,
    ) -> int:
        """
        Insert a section at the middle of the document.

        Args:
            heading: Section heading text
            paragraphs: List of paragraph texts
            heading_level: Heading level (1-9)

        Returns:
            Index where content was inserted
        """
        middle_index = len(self._doc.paragraphs) // 2
        return self.insert_section_after_paragraph(
            paragraph_index=middle_index,
            heading=heading,
            paragraphs=paragraphs,
            heading_level=heading_level,
        )

    def save(self) -> bytes:
        """
        Save the modified document to bytes.

        Returns:
            Modified document as bytes
        """
        buffer = io.BytesIO()
        self._doc.save(buffer)
        return buffer.getvalue()

    def save_to_file(self, file_path: str) -> None:
        """
        Save the modified document to a file.

        Args:
            file_path: Path to save the document
        """
        self._doc.save(file_path)


# Convenience function
async def import_document(
    file_path: str,
    doc_store: Optional[DocumentStore] = None,
) -> DocumentModel:
    """
    Convenience function to import a Word document.

    Args:
        file_path: Path to .docx file
        doc_store: Optional document store

    Returns:
        Imported DocumentModel
    """
    importer = DocumentImporter(doc_store=doc_store)
    await importer.connect()
    try:
        return await importer.import_from_file(file_path)
    finally:
        await importer.close()
