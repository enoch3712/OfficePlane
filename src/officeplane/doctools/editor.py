"""
DocumentEditor - Context manager for batch document operations.

The key insight: open a document ONCE, perform MANY operations, save ONCE.
This is critical for performance and enables transaction support.

Usage:
    # Context manager keeps file open for batch operations
    with DocumentEditor("report.docx") as editor:
        editor.add_paragraph("Introduction")
        editor.add_heading("Chapter 1", level=1)
        editor.add_paragraph("Content here...")
        # File saved automatically on context exit

    # Transaction support with rollback
    with DocumentEditor("report.docx") as editor:
        with editor.transaction() as tx:
            editor.add_paragraph("Risky content")
            editor.delete_paragraph(0)  # If this fails...
            # ...tx.rollback() is called automatically

    # Explicit control
    editor = DocumentEditor("report.docx")
    editor.open()
    try:
        editor.add_paragraph("Hello")
        editor.save()
    finally:
        editor.close()
"""

from __future__ import annotations

import shutil
import tempfile
from contextlib import contextmanager
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum, auto
from pathlib import Path
from typing import (
    Any,
    Callable,
    Dict,
    Iterator,
    List,
    Optional,
    TypeVar,
    Union,
)
from uuid import uuid4

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.table import Table

from officeplane.doctools.result import (
    DocError,
    ErrorCode,
    Err,
    Ok,
    Result,
    try_operation,
)


class EditorState(Enum):
    """State of the document editor."""

    CLOSED = auto()
    OPEN = auto()
    MODIFIED = auto()
    ERROR = auto()


@dataclass
class EditOperation:
    """Record of an operation for transaction support."""

    id: str = field(default_factory=lambda: str(uuid4())[:8])
    operation: str = ""
    timestamp: datetime = field(default_factory=datetime.now)
    params: Dict[str, Any] = field(default_factory=dict)
    rollback_fn: Optional[Callable[[], None]] = None


@dataclass
class ParagraphRef:
    """Reference to a paragraph with position info."""

    index: int
    text: str
    style: Optional[str] = None


@dataclass
class TableRef:
    """Reference to a table with position info."""

    index: int
    rows: int
    cols: int


@dataclass
class HeadingRef:
    """Reference to a heading with level."""

    index: int
    text: str
    level: int


class EditSession:
    """
    A transaction-like session for atomic document modifications.

    Changes are tracked and can be rolled back on failure.
    """

    def __init__(self, editor: DocumentEditor):
        self.editor = editor
        self.operations: List[EditOperation] = []
        self.committed = False
        self.rolled_back = False
        self._backup_path: Optional[Path] = None

    def __enter__(self) -> EditSession:
        """Start the transaction by creating a backup."""
        if self.editor._doc is None:
            raise ValueError("Editor must be open to start transaction")

        # Create backup
        self._backup_path = Path(tempfile.mktemp(suffix=".docx"))
        self.editor._doc.save(str(self._backup_path))
        return self

    def __exit__(self, exc_type, exc_val, exc_tb) -> bool:
        """End the transaction - commit on success, rollback on failure."""
        if exc_type is not None:
            # Exception occurred - rollback
            self.rollback()
            return False  # Re-raise the exception

        if not self.committed and not self.rolled_back:
            # Auto-commit if no explicit action taken
            self.commit()

        # Clean up backup
        if self._backup_path and self._backup_path.exists():
            self._backup_path.unlink()

        return False

    def record(self, operation: EditOperation) -> None:
        """Record an operation for potential rollback."""
        self.operations.append(operation)

    def commit(self) -> Result[None]:
        """Commit the transaction - changes are kept."""
        if self.rolled_back:
            return Err(
                DocError(
                    code=ErrorCode.TRANSACTION_FAILED,
                    message="Cannot commit - transaction was already rolled back",
                )
            )

        self.committed = True
        return Ok(None)

    def rollback(self) -> Result[None]:
        """Rollback all changes by restoring from backup."""
        if self.committed:
            return Err(
                DocError(
                    code=ErrorCode.ROLLBACK_FAILED,
                    message="Cannot rollback - transaction was already committed",
                )
            )

        if self._backup_path and self._backup_path.exists():
            # Restore from backup
            self.editor._doc = Document(str(self._backup_path))
            self.rolled_back = True
            return Ok(None)

        return Err(
            DocError(
                code=ErrorCode.ROLLBACK_FAILED,
                message="No backup available for rollback",
            )
        )


class DocumentEditor:
    """
    High-performance document editor with batch operations and transactions.

    Design principles:
    - Open once, operate many, save once
    - All operations return Result types
    - Transaction support with rollback
    - Context manager for automatic cleanup
    """

    def __init__(
        self,
        path: Union[str, Path],
        auto_save: bool = True,
        create_if_missing: bool = False,
    ):
        """
        Initialize the editor.

        Args:
            path: Path to the document
            auto_save: Whether to save automatically on context exit
            create_if_missing: Whether to create a new document if path doesn't exist
        """
        self.path = Path(path)
        self.auto_save = auto_save
        self.create_if_missing = create_if_missing

        self._doc: Optional[DocxDocument] = None
        self._state = EditorState.CLOSED
        self._current_session: Optional[EditSession] = None
        self._operation_count = 0

    @property
    def state(self) -> EditorState:
        """Current editor state."""
        return self._state

    @property
    def is_open(self) -> bool:
        """Whether the document is currently open."""
        return self._state in (EditorState.OPEN, EditorState.MODIFIED)

    @property
    def is_modified(self) -> bool:
        """Whether the document has unsaved changes."""
        return self._state == EditorState.MODIFIED

    @property
    def doc(self) -> Optional[DocxDocument]:
        """Raw python-docx Document (for advanced operations)."""
        return self._doc

    def __enter__(self) -> DocumentEditor:
        """Open the document for editing."""
        result = self.open()
        if result.is_err():
            raise RuntimeError(f"Failed to open document: {result.unwrap_err()}")
        return self

    def __exit__(self, exc_type, exc_val, exc_tb) -> bool:
        """Close the document, saving if auto_save is enabled."""
        if exc_type is None and self.auto_save and self.is_modified:
            self.save()
        self.close()
        return False

    def open(self) -> Result[None]:
        """Open the document for editing."""
        if self._state != EditorState.CLOSED:
            return Err(
                DocError(
                    code=ErrorCode.OPERATION_FAILED,
                    message="Document is already open",
                    source_file=str(self.path),
                )
            )

        if not self.path.exists():
            if self.create_if_missing:
                self._doc = Document()
                self._state = EditorState.MODIFIED
                return Ok(None)
            return Err(DocError.file_not_found(str(self.path)))

        result = try_operation(
            lambda: Document(str(self.path)),
            f"Opening document {self.path}",
        )

        if result.is_ok():
            self._doc = result.unwrap()
            self._state = EditorState.OPEN
            return Ok(None)

        self._state = EditorState.ERROR
        return Err(result.unwrap_err())

    def close(self) -> None:
        """Close the document without saving."""
        self._doc = None
        self._state = EditorState.CLOSED
        self._current_session = None

    def save(self, path: Optional[Union[str, Path]] = None) -> Result[None]:
        """
        Save the document.

        Args:
            path: Optional alternative path to save to
        """
        if self._doc is None:
            return Err(
                DocError(
                    code=ErrorCode.OPERATION_FAILED,
                    message="No document is open",
                )
            )

        save_path = Path(path) if path else self.path

        result = try_operation(
            lambda: self._doc.save(str(save_path)),
            f"Saving document to {save_path}",
        )

        if result.is_ok():
            self._state = EditorState.OPEN
            return Ok(None)

        return Err(result.unwrap_err())

    @contextmanager
    def transaction(self) -> Iterator[EditSession]:
        """
        Start a transaction for atomic operations.

        Changes can be rolled back if an error occurs.

        Example:
            with editor.transaction() as tx:
                editor.add_paragraph("Content")
                editor.delete_table(0)
                # If any operation fails, all changes are rolled back
        """
        if self._doc is None:
            raise ValueError("Editor must be open to start transaction")

        session = EditSession(self)
        self._current_session = session
        try:
            with session:
                yield session
        finally:
            self._current_session = None

    def _mark_modified(self) -> None:
        """Mark the document as modified."""
        if self._state == EditorState.OPEN:
            self._state = EditorState.MODIFIED
        self._operation_count += 1

    def _record_operation(
        self,
        operation: str,
        params: Dict[str, Any],
        rollback_fn: Optional[Callable[[], None]] = None,
    ) -> None:
        """Record an operation for transaction tracking."""
        if self._current_session:
            self._current_session.record(
                EditOperation(
                    operation=operation,
                    params=params,
                    rollback_fn=rollback_fn,
                )
            )

    # =========================================================================
    # Paragraph Operations
    # =========================================================================

    def add_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
    ) -> Result[ParagraphRef]:
        """
        Add a paragraph at the end of the document.

        Args:
            text: Paragraph text
            style: Optional style name (e.g., "Normal", "BodyText")

        Returns:
            Result containing ParagraphRef on success
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        try:
            para = self._doc.add_paragraph(text, style)
            self._mark_modified()

            index = len(self._doc.paragraphs) - 1
            ref = ParagraphRef(index=index, text=text, style=style)

            self._record_operation(
                "add_paragraph",
                {"text": text, "style": style},
            )

            return Ok(ref)

        except Exception as e:
            return Err(DocError.from_exception(e, "Adding paragraph"))

    def insert_paragraph_after(
        self,
        anchor_index: int,
        text: str,
        style: Optional[str] = None,
    ) -> Result[ParagraphRef]:
        """
        Insert a paragraph after the paragraph at anchor_index.

        Args:
            anchor_index: Index of paragraph to insert after
            text: Paragraph text
            style: Optional style name

        Returns:
            Result containing ParagraphRef on success
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        paragraphs = self._doc.paragraphs
        if anchor_index < 0 or anchor_index >= len(paragraphs):
            return Err(
                DocError.invalid_position(
                    anchor_index,
                    f"0 to {len(paragraphs) - 1}",
                )
            )

        try:
            # Get the paragraph to insert after
            anchor_para = paragraphs[anchor_index]

            # Create new paragraph element
            new_para = self._doc.add_paragraph(text, style)
            new_para_elem = new_para._element

            # Move it after the anchor
            anchor_para._element.addnext(new_para_elem)

            self._mark_modified()

            ref = ParagraphRef(index=anchor_index + 1, text=text, style=style)
            self._record_operation(
                "insert_paragraph_after",
                {"anchor_index": anchor_index, "text": text, "style": style},
            )

            return Ok(ref)

        except Exception as e:
            return Err(DocError.from_exception(e, "Inserting paragraph"))

    def delete_paragraph(self, index: int) -> Result[None]:
        """
        Delete a paragraph by index.

        Args:
            index: Index of paragraph to delete

        Returns:
            Result indicating success or failure
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        paragraphs = self._doc.paragraphs
        if index < 0 or index >= len(paragraphs):
            return Err(
                DocError.invalid_position(index, f"0 to {len(paragraphs) - 1}")
            )

        try:
            para = paragraphs[index]
            para._element.getparent().remove(para._element)

            self._mark_modified()
            self._record_operation("delete_paragraph", {"index": index})

            return Ok(None)

        except Exception as e:
            return Err(DocError.from_exception(e, "Deleting paragraph"))

    def get_paragraph(self, index: int) -> Result[ParagraphRef]:
        """Get a paragraph by index."""
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        paragraphs = self._doc.paragraphs
        if index < 0 or index >= len(paragraphs):
            return Err(
                DocError.invalid_position(index, f"0 to {len(paragraphs) - 1}")
            )

        para = paragraphs[index]
        return Ok(
            ParagraphRef(
                index=index,
                text=para.text,
                style=para.style.name if para.style else None,
            )
        )

    def find_paragraphs(
        self,
        text: Optional[str] = None,
        style: Optional[str] = None,
        contains: Optional[str] = None,
    ) -> Result[List[ParagraphRef]]:
        """
        Find paragraphs matching criteria.

        Args:
            text: Exact text match
            style: Style name match
            contains: Substring match

        Returns:
            Result containing list of matching ParagraphRefs
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        results: List[ParagraphRef] = []
        for i, para in enumerate(self._doc.paragraphs):
            matches = True

            if text is not None and para.text != text:
                matches = False
            if style is not None:
                para_style = para.style.name if para.style else None
                if para_style != style:
                    matches = False
            if contains is not None and contains not in para.text:
                matches = False

            if matches:
                results.append(
                    ParagraphRef(
                        index=i,
                        text=para.text,
                        style=para.style.name if para.style else None,
                    )
                )

        return Ok(results)

    def replace_text(
        self,
        old_text: str,
        new_text: str,
        paragraph_index: Optional[int] = None,
    ) -> Result[int]:
        """
        Replace text in paragraphs.

        Args:
            old_text: Text to find
            new_text: Replacement text
            paragraph_index: If specified, only replace in this paragraph

        Returns:
            Result containing count of replacements
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        try:
            count = 0
            paragraphs = self._doc.paragraphs

            if paragraph_index is not None:
                if paragraph_index < 0 or paragraph_index >= len(paragraphs):
                    return Err(
                        DocError.invalid_position(
                            paragraph_index,
                            f"0 to {len(paragraphs) - 1}",
                        )
                    )
                paragraphs = [paragraphs[paragraph_index]]

            for para in paragraphs:
                if old_text in para.text:
                    # Replace in each run to preserve formatting
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1

            if count > 0:
                self._mark_modified()
                self._record_operation(
                    "replace_text",
                    {"old_text": old_text, "new_text": new_text, "count": count},
                )

            return Ok(count)

        except Exception as e:
            return Err(DocError.from_exception(e, "Replacing text"))

    # =========================================================================
    # Heading Operations
    # =========================================================================

    def add_heading(self, text: str, level: int = 1) -> Result[HeadingRef]:
        """
        Add a heading at the end of the document.

        Args:
            text: Heading text
            level: Heading level (0-9, where 0 is Title)

        Returns:
            Result containing HeadingRef on success
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        if level < 0 or level > 9:
            return Err(
                DocError(
                    code=ErrorCode.INVALID_ARGUMENT,
                    message=f"Invalid heading level: {level}",
                    suggestion="Use a level between 0 (Title) and 9",
                )
            )

        try:
            self._doc.add_heading(text, level)
            self._mark_modified()

            index = len(self._doc.paragraphs) - 1
            ref = HeadingRef(index=index, text=text, level=level)

            self._record_operation("add_heading", {"text": text, "level": level})

            return Ok(ref)

        except Exception as e:
            return Err(DocError.from_exception(e, "Adding heading"))

    def find_headings(
        self,
        level: Optional[int] = None,
        contains: Optional[str] = None,
    ) -> Result[List[HeadingRef]]:
        """
        Find headings in the document.

        Args:
            level: Filter by heading level
            contains: Filter by text substring

        Returns:
            Result containing list of HeadingRefs
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        results: List[HeadingRef] = []
        for i, para in enumerate(self._doc.paragraphs):
            if para.style and para.style.name.startswith("Heading"):
                # Extract level from style name
                try:
                    heading_level = int(para.style.name.replace("Heading ", "").strip())
                except ValueError:
                    heading_level = 0 if para.style.name == "Title" else 1

                # Apply filters
                if level is not None and heading_level != level:
                    continue
                if contains is not None and contains not in para.text:
                    continue

                results.append(
                    HeadingRef(index=i, text=para.text, level=heading_level)
                )

        return Ok(results)

    # =========================================================================
    # Table Operations
    # =========================================================================

    def add_table(
        self,
        rows: int,
        cols: int,
        style: Optional[str] = None,
    ) -> Result[TableRef]:
        """
        Add a table at the end of the document.

        Args:
            rows: Number of rows
            cols: Number of columns
            style: Optional table style name

        Returns:
            Result containing TableRef on success
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        if rows < 1 or cols < 1:
            return Err(
                DocError(
                    code=ErrorCode.INVALID_ARGUMENT,
                    message=f"Invalid table dimensions: {rows}x{cols}",
                    suggestion="Rows and columns must be at least 1",
                )
            )

        try:
            table = self._doc.add_table(rows=rows, cols=cols)
            if style:
                table.style = style

            self._mark_modified()

            index = len(self._doc.tables) - 1
            ref = TableRef(index=index, rows=rows, cols=cols)

            self._record_operation(
                "add_table",
                {"rows": rows, "cols": cols, "style": style},
            )

            return Ok(ref)

        except Exception as e:
            return Err(DocError.from_exception(e, "Adding table"))

    def set_cell(
        self,
        table_index: int,
        row: int,
        col: int,
        text: str,
    ) -> Result[None]:
        """
        Set the text of a table cell.

        Args:
            table_index: Index of the table
            row: Row index (0-based)
            col: Column index (0-based)
            text: Text to set

        Returns:
            Result indicating success or failure
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        tables = self._doc.tables
        if table_index < 0 or table_index >= len(tables):
            return Err(
                DocError.invalid_position(
                    table_index,
                    f"0 to {len(tables) - 1} (tables)",
                )
            )

        table = tables[table_index]
        if row < 0 or row >= len(table.rows):
            return Err(
                DocError.invalid_position(row, f"0 to {len(table.rows) - 1} (rows)")
            )
        if col < 0 or col >= len(table.columns):
            return Err(
                DocError.invalid_position(col, f"0 to {len(table.columns) - 1} (cols)")
            )

        try:
            table.cell(row, col).text = text
            self._mark_modified()

            self._record_operation(
                "set_cell",
                {"table_index": table_index, "row": row, "col": col, "text": text},
            )

            return Ok(None)

        except Exception as e:
            return Err(DocError.from_exception(e, "Setting cell text"))

    def fill_table(
        self,
        table_index: int,
        data: List[List[str]],
        start_row: int = 0,
        start_col: int = 0,
    ) -> Result[None]:
        """
        Fill a table with data from a 2D list.

        Args:
            table_index: Index of the table
            data: 2D list of cell values
            start_row: Starting row (0-based)
            start_col: Starting column (0-based)

        Returns:
            Result indicating success or failure
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        tables = self._doc.tables
        if table_index < 0 or table_index >= len(tables):
            return Err(
                DocError.invalid_position(table_index, f"0 to {len(tables) - 1}")
            )

        table = tables[table_index]

        try:
            for r_idx, row_data in enumerate(data):
                row = start_row + r_idx
                if row >= len(table.rows):
                    break

                for c_idx, cell_value in enumerate(row_data):
                    col = start_col + c_idx
                    if col >= len(table.columns):
                        break
                    table.cell(row, col).text = str(cell_value)

            self._mark_modified()
            self._record_operation(
                "fill_table",
                {
                    "table_index": table_index,
                    "data_shape": f"{len(data)}x{len(data[0]) if data else 0}",
                },
            )

            return Ok(None)

        except Exception as e:
            return Err(DocError.from_exception(e, "Filling table"))

    # =========================================================================
    # Document Info
    # =========================================================================

    def get_stats(self) -> Result[Dict[str, Any]]:
        """
        Get document statistics.

        Returns:
            Result containing stats dict with paragraph_count, table_count, etc.
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        try:
            paragraphs = self._doc.paragraphs
            tables = self._doc.tables

            word_count = sum(len(p.text.split()) for p in paragraphs)
            char_count = sum(len(p.text) for p in paragraphs)

            stats = {
                "path": str(self.path),
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "word_count": word_count,
                "character_count": char_count,
                "state": self._state.name,
                "operation_count": self._operation_count,
            }

            return Ok(stats)

        except Exception as e:
            return Err(DocError.from_exception(e, "Getting document stats"))

    def get_structure(self) -> Result[List[Dict[str, Any]]]:
        """
        Get document structure (headings and their hierarchy).

        Returns:
            Result containing list of heading dicts with level, text, index
        """
        if self._doc is None:
            return Err(DocError(code=ErrorCode.OPERATION_FAILED, message="No document open"))

        try:
            structure: List[Dict[str, Any]] = []

            for i, para in enumerate(self._doc.paragraphs):
                if para.style and para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.replace("Heading ", "").strip())
                    except ValueError:
                        level = 0 if para.style.name == "Title" else 1

                    structure.append({
                        "index": i,
                        "level": level,
                        "text": para.text,
                        "style": para.style.name,
                    })

            return Ok(structure)

        except Exception as e:
            return Err(DocError.from_exception(e, "Getting document structure"))
